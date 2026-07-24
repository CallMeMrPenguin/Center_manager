import os
import sys
import json
import math
from typing import List, Dict, Any

# python-docx imports
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Local instruction mapping dictionary
INSTRUCTION_MAP = {
    "pr": "Mark the letter A, B, C, or D on your answer sheet to indicate the word whose underlined part differs from the other three in pronunciation in each of the following questions.",
    "mq": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct answer to each of the following questions.",
    "sg": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct meaning of the sign in each of the following questions.",
    "nt": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct meaning of the notice in each of the following questions.",
    "cz": "Read the following passage and mark the letter A, B, C, or D on your answer sheet to indicate the correct word or phrase that best fits each of the numbered blanks.",
    "ro": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct arrangement of the sentences to make a meaningful text in each of the following questions.",
    "rd": "Read the following passage and mark the letter A, B, C, or D on your answer sheet to indicate the correct answer to each of the following questions.",
    "er": "Mark the letter A, B, C, or D on your answer sheet to indicate the underlined part that needs correction in each of the following questions."
}

def parse_text_formatting(text: str) -> List[Dict[str, Any]]:
    import re
    if not text:
        return []
        
    pattern = re.compile(
        r'(?P<err>\[(?P<err_txt>[^\]]+)\]\((?P<err_let>[A-D])\))|'
        r'(?P<bold_under1>\*\*\[(?P<bu_txt1>[^\]]+)\]\*\*)|'
        r'(?P<bold_under2>\[\*\*(?P<bu_txt2>[^*]+)\*\*\])|'
        r'(?P<bold>\*\*(?P<b_txt>[^*]+)\*\*)|'
        r'(?P<italic>\*(?P<i_txt>[^*]+)\*)|'
        r'(?P<under>\[(?P<u_txt>[^\]]+)\])'
    )
    
    segments = []
    last_idx = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            segments.append({
                "text": text[last_idx:start],
                "bold": False,
                "italic": False,
                "underline": False
            })
        
        gd = match.groupdict()
        if gd['err']:
            segments.append({"text": gd['err_txt'], "bold": False, "italic": False, "underline": True})
            segments.append({"text": f" ({gd['err_let']})", "bold": True, "italic": False, "underline": False})
        elif gd['bold_under1'] or gd['bold_under2']:
            txt = gd['bu_txt1'] if gd['bold_under1'] else gd['bu_txt2']
            segments.append({"text": txt, "bold": True, "italic": False, "underline": True})
        elif gd['bold']:
            segments.append({"text": gd['b_txt'], "bold": True, "italic": False, "underline": False})
        elif gd['italic']:
            segments.append({"text": gd['i_txt'], "bold": False, "italic": True, "underline": False})
        elif gd['under']:
            segments.append({"text": gd['u_txt'], "bold": False, "italic": False, "underline": True})
            
        last_idx = end
        
    if last_idx < len(text):
        segments.append({
            "text": text[last_idx:],
            "bold": False,
            "italic": False,
            "underline": False
        })
        
    return segments


class WordDocumentCompiler:
    def __init__(self, settings: Dict[str, Any] = None):
        self.settings = settings or {}
        self.doc = Document()
        self._configure_page()
        self._configure_styles()

    def _configure_page(self):
        top_cm = self.settings.get("margin_top", 2.0)
        bottom_cm = self.settings.get("margin_bottom", 2.0)
        left_cm = self.settings.get("margin_left", 3.0)
        right_cm = self.settings.get("margin_right", 1.5)
        
        for section in self.doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(top_cm)
            section.bottom_margin = Cm(bottom_cm)
            section.left_margin = Cm(left_cm)
            section.right_margin = Cm(right_cm)

    def _configure_styles(self):
        font_name = self.settings.get("font_name", "Times New Roman")
        font_size = self.settings.get("font_size", 12.0)
        line_spacing = self.settings.get("line_spacing", 1.15)
        space_after = self.settings.get("space_after", 6.0)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        p_format = style.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(space_after)
        p_format.line_spacing = line_spacing

    def add_instruction_header(self, text: str):
        space_before = self.settings.get("header_space_before", 14.0)
        space_after = self.settings.get("header_space_after", 8.0)
        font_size = self.settings.get("font_size", 12.0)
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        # Directions instruction text is bold directly (no Parts prefixes)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)

    def add_question(self, q_num: Any, q_text: str):
        space_before = self.settings.get("question_space_before", 6.0)
        space_after = self.settings.get("question_space_after", 4.0)
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run_q = p.add_run(f"Question {q_num}: ")
        run_q.bold = True
        
        segments = parse_text_formatting(q_text)
        for seg in segments:
            run = p.add_run(seg["text"])
            if seg["bold"]:
                run.bold = True
            if seg["italic"]:
                run.italic = True
            if seg["underline"]:
                run.underline = True

    def add_options_grid(self, options: List[str], exercise_type: str):
        if not options:
            return
            
        max_len = max(len(str(opt)) for opt in options) if options else 0
        
        if exercise_type in ["pr", "st"] or max_len < 15:
            cols = 4
        elif max_len < 35:
            cols = 2
        else:
            cols = 1

        left_indent_cm = self.settings.get("options_left_indent", 0.5)
        space_before = self.settings.get("options_space_before", 0.0)
        space_after = self.settings.get("options_space_after", 3.0)
        
        # Calculate column positions based on left/right margins and indents
        left_margin_cm = self.settings.get("margin_left", 3.0)
        right_margin_cm = self.settings.get("margin_right", 1.5)
        printable_width_cm = 21.0 - left_margin_cm - right_margin_cm
        remaining_width_cm = printable_width_cm - left_indent_cm

        if cols == 1:
            for idx, opt in enumerate(options):
                prefix = chr(65 + idx)
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(left_indent_cm)  # Customized left indent (5mm by default)
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                
                run_prefix = p.add_run(f"{prefix}. ")
                run_prefix.bold = True
                
                segments = parse_text_formatting(opt)
                for seg in segments:
                    run = p.add_run(seg["text"])
                    if seg["bold"]: run.bold = True
                    if seg["italic"]: run.italic = True
                    if seg["underline"]: run.underline = True
                    
        elif cols == 2:
            # 2 rows, 2 options per row (A & B, C & D)
            col_width = remaining_width_cm / 2
            tab_pos = left_indent_cm + col_width
            
            for row in range(2):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(left_indent_cm)
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_pos), WD_TAB_ALIGNMENT.LEFT)
                
                # Option 1 (A or C)
                idx1 = row * 2
                if idx1 < len(options):
                    prefix = chr(65 + idx1)
                    run_prefix = p.add_run(f"{prefix}. ")
                    run_prefix.bold = True
                    for seg in parse_text_formatting(options[idx1]):
                        run = p.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                        
                # Tab spacing
                p.add_run("\t")
                
                # Option 2 (B or D)
                idx2 = row * 2 + 1
                if idx2 < len(options):
                    prefix = chr(65 + idx2)
                    run_prefix = p.add_run(f"{prefix}. ")
                    run_prefix.bold = True
                    for seg in parse_text_formatting(options[idx2]):
                        run = p.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                        
        elif cols == 4:
            # 1 row containing all 4 options (A, B, C, D)
            col_width = remaining_width_cm / 4
            tab1 = left_indent_cm + col_width
            tab2 = left_indent_cm + (col_width * 2)
            tab3 = left_indent_cm + (col_width * 3)
            
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(left_indent_cm)
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(tab1), WD_TAB_ALIGNMENT.LEFT)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(tab2), WD_TAB_ALIGNMENT.LEFT)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(tab3), WD_TAB_ALIGNMENT.LEFT)
            
            for idx in range(4):
                if idx < len(options):
                    prefix = chr(65 + idx)
                    run_prefix = p.add_run(f"{prefix}. ")
                    run_prefix.bold = True
                    for seg in parse_text_formatting(options[idx]):
                        run = p.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                
                # Add tab character to advance to next column
                if idx < 3:
                    p.add_run("\t")

    def compile(self, exercises: List[Dict[str, Any]], output_filepath: Any):
        # Handle dictionary wrapper if present (e.g. {"defs": ..., "data": [...]})
        if isinstance(exercises, dict):
            for key in ["data", "exercises", "questions", "list"]:
                if key in exercises and isinstance(exercises[key], list):
                    exercises = exercises[key]
                    break
            else:
                if "t" in exercises or "q" in exercises:
                    exercises = [exercises]
                else:
                    raise ValueError("JSON dictionary does not contain a list of exercises.")

        current_type = None
        for ex in exercises:
            ex_type = ex.get("t")
            
            # Whenever the exercise type shifts, inject the instruction block
            if ex_type != current_type:
                current_type = ex_type
                instruction_text = INSTRUCTION_MAP.get(current_type)
                if instruction_text:
                    self.add_instruction_header(instruction_text)
            
            # Print logic based on exercise type structure
            if ex_type in ["cz", "rd"]:
                # Print the passage paragraphs
                space_before = self.settings.get("passage_space_before", 4.0)
                space_after = self.settings.get("passage_space_after", 6.0)
                indent_first = self.settings.get("passage_indent_first", 0.75)
                
                for paragraph_text in ex.get("b", []):
                    p_passage = self.doc.add_paragraph()
                    p_passage.paragraph_format.space_before = Pt(space_before)
                    p_passage.paragraph_format.space_after = Pt(space_after)
                    p_passage.paragraph_format.first_line_indent = Cm(indent_first)
                    
                    segments = parse_text_formatting(paragraph_text)
                    for seg in segments:
                        run = p_passage.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                
                # Print passage sub-questions
                for sub in ex.get("k", []):
                    sub_q = sub.get("q")
                    sub_x = sub.get("x", "")
                    sub_o = sub.get("o", [])
                    
                    self.add_question(sub_q, sub_x)
                    self.add_options_grid(sub_o, ex_type)
                    
            elif ex_type == "ro":
                # Sentence reordering
                q_num = ex.get("q")
                q_text = ex.get("x", "Choose the best arrangement of the sentences:")
                self.add_question(q_num, q_text)
                
                space_before = self.settings.get("reorder_space_before", 0.0)
                space_after = self.settings.get("reorder_space_after", 2.0)
                left_indent = self.settings.get("reorder_left_indent", 1.0)
                
                # Print sentences to reorder
                for item in ex.get("i", []):
                    p_item = self.doc.add_paragraph()
                    p_item.paragraph_format.left_indent = Cm(left_indent)
                    p_item.paragraph_format.space_before = Pt(space_before)
                    p_item.paragraph_format.space_after = Pt(space_after)
                    
                    segments = parse_text_formatting(item)
                    for seg in segments:
                        run = p_item.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                
                # Print option grid layout
                self.add_options_grid(ex.get("o", []), ex_type)
                
            elif ex_type == "nt":
                # Notice question card
                q_num = ex.get("q")
                q_text = ex.get("x", "")
                self.add_question(q_num, q_text)
                
                space_before = self.settings.get("notice_space_before", 4.0)
                space_after = self.settings.get("notice_space_after", 6.0)
                left_indent = self.settings.get("notice_left_indent", 1.0)
                
                # Print notice text block
                p_body = self.doc.add_paragraph()
                p_body.paragraph_format.left_indent = Cm(left_indent)
                p_body.paragraph_format.space_before = Pt(space_before)
                p_body.paragraph_format.space_after = Pt(space_after)
                
                segments = parse_text_formatting(ex.get("b", ""))
                for seg in segments:
                    run = p_body.add_run(seg["text"])
                    run.italic = True
                    if seg["bold"]: run.bold = True
                    if seg["underline"]: run.underline = True
                
                # Print options grid layout
                self.add_options_grid(ex.get("o", []), ex_type)
                
            else:
                # Flat types: pr, mq, sg, er, etc.
                q_num = ex.get("q")
                q_text = ex.get("x", "")
                options = ex.get("o", [])
                
                self.add_question(q_num, q_text)
                self.add_options_grid(options, ex_type)
        
        # Save to filepath or buffer stream
        self.doc.save(output_filepath)


MOCK_PAYLOAD = [
    {
        "t": "pr",
        "q": 1,
        "x": "",
        "o": ["pass[ed]", "plann[ed]", "hopp[ed]", "play[ed]"],
        "a": "D"
    },
    {
        "t": "st",
        "q": 2,
        "x": "",
        "o": ["teacher", "student", "decide", "member"],
        "a": "C"
    },
    {
        "t": "mq",
        "q": 3,
        "x": "If it **rains** tomorrow, we *will cancel* the picnic.",
        "o": ["cancel", "would cancel", "will cancel", "canceled"],
        "a": "C"
    },
    {
        "t": "er",
        "q": 4,
        "x": "Because of [his](A) illness, he [could not](B) go to school, [so](C) he was [sadly](D).",
        "o": ["his", "could not", "so", "sadly"],
        "a": "D"
    }
]


def main():
    if len(sys.argv) < 3:
        print("Usage: python compiler.py <input_json_path> <output_docx_path>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(input_path):
        print(f"Error: Input file {input_path} does not exist.")
        sys.exit(1)
        
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            exercises = json.load(f)
            
        compiler = WordDocumentCompiler()
        compiler.compile(exercises, output_path)
        print(f"SUCCESS: Document compiled to {output_path}")
        sys.exit(0)
    except Exception as e:
        print(f"ERROR: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
