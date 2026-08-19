import os
import re
import shutil
import copy
import docx
import win32com.client
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Unit Catalog for Global Success (Full Year: Unit 1-12)
UNIT_NAMES = {
    6: {
        1: "MY NEW SCHOOL",
        2: "MY HOUSE",
        3: "MY FRIENDS",
        4: "MY NEIGHBOURHOOD",
        5: "NATURAL WONDERS OF THE VIET NAM",
        6: "OUR TET HOLIDAY",
        7: "TELEVISION",
        8: "SPORTS AND GAMES",
        9: "CITIES OF THE WORLD",
        10: "OUR HOUSES IN THE FUTURE",
        11: "OUR GREENER WORLD",
        12: "ROBOTS"
    },
    7: {
        1: "HOBBIES",
        2: "HEALTHY LIVING",
        3: "COMMUNITY SERVICE",
        4: "MUSIC AND ARTS",
        5: "FOOD AND DRINK",
        6: "A VISIT TO SCHOOL",
        7: "TRAFFIC",
        8: "FILMS",
        9: "FESTIVALS AROUND THE WORLD",
        10: "ENERGY SOURCES",
        11: "TRAVELLING IN THE FUTURE",
        12: "ENGLISH-SPEAKING COUNTRIES"
    },
    8: {
        1: "LEISURE TIME",
        2: "LIFE IN THE COUNTRYSIDE",
        3: "TEENAGERS",
        4: "ETHNIC GROUPS OF VIET NAM",
        5: "OUR CUSTOMS AND TRADITIONS",
        6: "LIFESTYLES",
        7: "ENVIRONMENTAL PROTECTION",
        8: "SHOPPING",
        9: "NATURAL DISASTERS",
        10: "COMMUNICATION IN THE FUTURE",
        11: "SCIENCE AND TECHNOLOGY",
        12: "LIFE ON OTHER PLANETS"
    },
    9: {
        1: "LOCAL COMMUNITY",
        2: "CITY LIFE",
        3: "TEEN STRESS AND PRESSURE",
        4: "LIFE IN THE PAST",
        5: "WONDERS OF VIET NAM",
        6: "VIET NAM THEN AND NOW",
        7: "NATURAL WONDERS OF THE WORLD",
        8: "TOURISM",
        9: "WORLD ENGLISHES",
        10: "PLANET EARTH",
        11: "ELECTRONIC DEVICES",
        12: "CAREER PATHS"
    }
}

def append_document(doc_target, doc_source):
    """Appends all body elements of doc_source to the end of doc_target, preserving formatting."""
    body_source = doc_source.element.body
    body_target = doc_target.element.body
    
    for element in body_source:
        # Ignore the final sectPr element of the source document
        if element.tag.endswith('sectPr'):
            continue
        new_element = copy.deepcopy(element)
        body_target.append(new_element)

def make_section_landscape(section):
    """Converts a section to Landscape and sets 1-inch margins."""
    new_width, new_height = section.page_height, section.page_width
    if section.orientation != WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_table_width_to_100(table):
    """Sets the table width to automatically stretch to 100% of the page width."""
    table.allow_autofit = True
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000') # 5000 pct = 100%
    tblW.set(qn('w:type'), 'pct')

def set_table_margins(table):
    """Sets cell margins to 1mm (approx 57 twips/dxa) on all sides."""
    tblPr = table._tbl.tblPr
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
        
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), '57') # 1mm ~ 57 dxa
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def repeat_table_header(table):
    """Sets the table header row to repeat on new pages and prevents splitting."""
    if len(table.rows) > 0:
        header_row = table.rows[0]
        trPr = header_row._tr.get_or_add_trPr()
        
        # Configure tblHeader
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            if qn('w:val') in tblHeader.attrib:
                del tblHeader.attrib[qn('w:val')]
                
        # Configure cantSplit
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)
        else:
            if qn('w:val') in cantSplit.attrib:
                del cantSplit.attrib[qn('w:val')]

def add_page_number_to_run(run):
    """Inserts XML fields to dynamically calculate and display page numbers in Word."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_toc_to_paragraph(p):
    """Inserts a TOC (Table of Contents) field code into a paragraph."""
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    p.add_run("Updating Table of Contents...")
    run3 = p.add_run()
    run3._r.append(fldChar3)

def insert_section_break_before(doc, target_p):
    """Inserts a Section Break (Next Page) right before the specified paragraph."""
    new_p_element = OxmlElement('w:p')
    target_p._p.getparent().insert(target_p._p.getparent().index(target_p._p), new_p_element)
    
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_element, doc)
    
    pPr = new_p_element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    pPr.append(sectPr)
    
    return new_p

def format_title_paragraph(p, text):
    """Formats unit title paragraphs as Heading 1, size 16, Times New Roman, bold."""
    p.text = ""
    p.style = 'Heading 1'
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

def format_toc_title_paragraph(p, text):
    """Formats TOC title paragraph as size 16, Times New Roman, bold (Normal style to avoid TOC inclusion)."""
    p.text = ""
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

def apply_vocabulary_formatting(p, is_single_para):
    """Applies bold/italic style configurations and removes brackets."""
    text = p.text.strip()
    if not text:
        return
    
    p.text = ""
    
    # Parse prefix (e.g. `(in-)` or `(un-)`) and base word
    prefix_match = re.match(r'^(\([^)]+-\))\s*(.*)$', text)
    if prefix_match:
        prefix = prefix_match.group(1)
        base_word = prefix_match.group(2).strip()
    else:
        prefix = None
        base_word = text
        
    bold = False
    italic = False
    display_word = base_word
    
    if base_word.startswith('[') and base_word.endswith(']'):
        bold = True
        italic = False
        display_word = base_word[1:-1]
    elif base_word.startswith('{') and base_word.endswith('}'):
        bold = False
        italic = True
        display_word = base_word[1:-1]
    else:
        if is_single_para:
            bold = True
            italic = False
        else:
            bold = False
            italic = False
            
    # Add prefix run if present
    if prefix:
        r_prefix = p.add_run(prefix + " ")
        r_prefix.bold = bold
        r_prefix.italic = italic
        r_prefix.font.name = 'Times New Roman'
        r_prefix.font.size = Pt(12)
        
    # Add main word run
    r_word = p.add_run(display_word)
    r_word.bold = bold
    r_word.italic = italic
    r_word.font.name = 'Times New Roman'
    r_word.font.size = Pt(12)

def update_docx_fields(filepath):
    """Opens MS Word via COM to update Table of Contents and page numbers automatically."""
    if not win32_available:
        print("  [win32com] win32com is not available. Skipping automatic field updates.")
        return
    abs_path = os.path.abspath(filepath)
    print(f"  [win32com] Opening Word to update fields for: {abs_path}")
    word = None
    try:
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except Exception:
            pass
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True  # Real-time visual output in MS Word
        word.DisplayAlerts = 0
        doc = word.Documents.Open(abs_path)
        
        # Update fields in the document
        doc.Fields.Update()
        
        # Update tables of contents
        for toc in doc.TablesOfContents:
            toc.Update()
            
        doc.Save()
        doc.Close()
        print("  [win32com] Successfully updated fields.")
    except Exception as e:
        print(f"  [win32com] Error updating fields: {e}")
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            if pythoncom:
                pythoncom.CoUninitialize()
        except Exception:
            pass

def process_grade(grade):
    """Combines Semester 1 & 2 files for a grade, then applies all formatting rules."""
    src1 = f"Vocabulary List - Grade {grade} - 1.docx"
    src2 = f"Vocabulary List - Grade {grade} - 2.docx"
    dest = f"GS {grade} _ Vocabulary.docx"
    
    print(f"--- Processing Grade {grade} ---")
    if not os.path.exists(src1) or not os.path.exists(src2):
        print(f"  Error: Source files '{src1}' or '{src2}' not found!")
        return
        
    print(f"  Merging: {src1} + {src2} -> {dest}")
    
    # Load documents
    doc_target = docx.Document(src1)
    doc_source = docx.Document(src2)
    
    # Append doc_source to doc_target
    append_document(doc_target, doc_source)
    
    # Save target temporarily
    doc_target.save(dest)
    
    # Load combined document
    doc = docx.Document(dest)
    
    # 0. Find and delete paragraph matching 'Vocabulary List - Grade ...'
    paras_to_remove = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.search(r'Vocabulary\s+List\s*-\s*Grade', text, re.IGNORECASE):
            paras_to_remove.append(p)
            print(f"  Found header line to remove: '{text}'")
            
    for p in paras_to_remove:
        try:
            p._p.getparent().remove(p._p)
        except Exception:
            pass

    # 1. Find unit titles and rename them
    unit_1_p = None
    for p in doc.paragraphs:
        text = p.text.strip()
        match = re.match(r'^unit\s+(\d+)', text, re.IGNORECASE)
        if match:
            unit_num = int(match.group(1))
            if grade in UNIT_NAMES and unit_num in UNIT_NAMES[grade]:
                topic_name = UNIT_NAMES[grade][unit_num]
                new_title = f"UNIT {unit_num}: {topic_name}"
                format_title_paragraph(p, new_title)
                p.paragraph_format.page_break_before = True
                print(f"  Renamed Unit {unit_num} -> '{new_title}'")
                
                if unit_num == 1:
                    unit_1_p = p
                    p.paragraph_format.page_break_before = False

    # 2. Add Table of Contents at the beginning
    if len(doc.paragraphs) > 0:
        first_p = doc.paragraphs[0]
        
        # TOC Title (Not Heading 1 style to avoid self-inclusion in TOC)
        toc_title_p = doc.add_paragraph()
        format_toc_title_paragraph(toc_title_p, "MỤC LỤC")
        toc_title_p.paragraph_format.page_break_before = False
        first_p._p.getparent().insert(first_p._p.getparent().index(first_p._p), toc_title_p._p)
        
        # TOC Field
        toc_p = doc.add_paragraph()
        add_toc_to_paragraph(toc_p)
        first_p._p.getparent().insert(first_p._p.getparent().index(first_p._p), toc_p._p)
        
        # TOC Note Paragraph
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_before = Pt(12)
        note_p.paragraph_format.space_after = Pt(6)
        note_run = note_p.add_run(
            "*(Lưu ý: Từ bôi đậm là từ thuộc chủ đề nên ghi; từ in nghiêng là word family cân nhắc ghi; "
            "từ không bôi đậm hay in nghiêng là từ cho lớp Kid .1)*"
        )
        note_run.italic = True
        note_run.font.name = 'Times New Roman'
        note_run.font.size = Pt(11)
        first_p._p.getparent().insert(first_p._p.getparent().index(first_p._p), note_p._p)
        
        # Spacing Paragraph
        space_p = doc.add_paragraph()
        first_p._p.getparent().insert(first_p._p.getparent().index(first_p._p), space_p._p)

    # 3. Insert Section Break before Unit 1
    if unit_1_p:
        insert_section_break_before(doc, unit_1_p)
        print("  Inserted section break before Unit 1")
    else:
        print("  WARNING: Unit 1 heading not found!")

    doc.save(dest)
    
    # Reload to parse the newly inserted section breaks
    doc = docx.Document(dest)
    
    # 4. Set orientation of all sections to Landscape
    for idx, section in enumerate(doc.sections):
        make_section_landscape(section)
        print(f"  Converted Section {idx} to Landscape")

    # 5. Configure page numbering starting from 1 for Section 1 (Unit 1 onwards)
    if len(doc.sections) > 1:
        sec2 = doc.sections[1]
        sec2.footer.is_linked_to_previous = False
        
        footer_p = sec2.footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Insert page number only
        run_num = footer_p.add_run()
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(11)
        add_page_number_to_run(run_num)
        
        # Configure start page number = 1
        sectPr_sec2 = sec2._sectPr
        pgNumType = sectPr_sec2.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr_sec2.append(pgNumType)
        pgNumType.set(qn('w:start'), '1')
        print("  Configured page numbers starting from 1 (centered) for Section 1")

    # 6. Format vocabulary tables
    print(f"  Formatting {len(doc.tables)} vocabulary tables...")
    for t_idx, table in enumerate(doc.tables):
        set_table_width_to_100(table)
        set_table_margins(table)
        repeat_table_header(table)
        
        for r_idx, row in enumerate(table.rows):
            # Format cell spacing for table content rows
            if r_idx > 0:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.0
            
            if r_idx == 0:
                continue
            
            if len(row.cells) > 1:
                vocab_cell = row.cells[1]
                paras = vocab_cell.paragraphs
                is_single_para = (len(paras) == 1)
                for p in paras:
                    apply_vocabulary_formatting(p, is_single_para)
                    
    doc.save(dest)
    
    # 7. Update Table of Contents dynamically using Word COM
    update_docx_fields(dest)
    print(f"Finished processing Grade {grade} successfully!\n")

def main():
    cwd = os.getcwd()
    print(f"Working directory: {cwd}")
    
    # Process Grades 6, 7, 8, and 9
    for grade in [6, 7, 8, 9]:
        process_grade(grade)
        
    print("All grades processed successfully!")

if __name__ == "__main__":
    main()
