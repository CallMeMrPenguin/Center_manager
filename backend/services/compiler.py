import os
import sys
import json
import math
import copy
import time
from typing import List, Dict, Any

# pywin32 imports
try:
    import pythoncom
    import win32com.client
    win32com_available = True
except Exception:
    win32com_available = False

# python-docx imports (fallback & parser support)
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Local instruction mapping dictionary
INSTRUCTION_MAP_MCQ = {
    "pr": "Mark the letter A, B, C, or D on your answer sheet to indicate the word whose underlined part differs from the other three in pronunciation in each of the following questions.",
    "st": "Mark the letter A, B, C, or D on your answer sheet to indicate the word that differs from the other three in the position of primary stress in each of the following questions.",
    "sy": "Mark the letter A, B, C, or D on your answer sheet to indicate the word(s) CLOSEST in meaning to the underlined word(s) in each of the following questions.",
    "an": "Mark the letter A, B, C, or D on your answer sheet to indicate the word(s) OPPOSITE in meaning to the underlined word(s) in each of the following questions.",
    "sg": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct meaning of the sign in each of the following questions.",
    "nt": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct meaning of the notice in each of the following questions.",
    "cz": "Read the following passage and mark the letter A, B, C, or D on your answer sheet to indicate the word or phrase that best fits each of the numbered blanks.",
    "ro": "Mark the letter A, B, C, or D on your answer sheet to indicate the answer that best fits each of the following questions.",
    "rd": "Read the following passage and mark the letter A, B, C, or D on your answer sheet to indicate the answer that best fits each of the following questions.",
    "er": "Mark the letter A, B, C, or D on your answer sheet to indicate the underlined part that needs correction in each of the following questions.",
    "fb": "Mark the letter A, B, C, or D on your answer sheet to indicate the word or phrase that best fits each blank in the following questions.",
    "rw": "Mark the letter A, B, C, or D on your answer sheet to indicate the sentence that is closest in meaning to each of the following questions.",
    "mq": "Mark the letter A, B, C, or D on your answer sheet to indicate the answer that best fits each of the following questions.",
    "wb": "Complete each blank with ONE suitable word from the box.",
    "wq": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct question for the underlined part in each of the following questions.",
    "mt": "Mark the letter A, B, C, or D on your answer sheet to indicate the correct match for each of the following questions.",
    "tf": "Mark the letter A, B, C, or D on your answer sheet to indicate whether each of the following statements is True or False."
}

INSTRUCTION_MAP_NON_MCQ = {
    "pr": "Write the word with the underlined sound for each of the following questions.",
    "st": "Write the position of primary stress for each of the following words.",
    "sy": "Provide a word or phrase closest in meaning to the underlined part in each of the following sentences.",
    "an": "Provide a word or phrase opposite in meaning to the underlined part in each of the following sentences.",
    "sg": "State the meaning of each of the following signs.",
    "nt": "State the meaning of each of the following notices.",
    "cz": "Read the following passage and fill in each numbered blank with ONE suitable word.",
    "ro": "Rearrange the given words or phrases to make complete, meaningful sentences.",
    "rd": "Read the following passage and answer the questions that follow.",
    "er": "Identify and correct the mistake in each of the following sentences.",
    "fb": "Complete each of the following sentences with the correct form of the word in brackets or a suitable word.",
    "rw": "Rewrite each of the following sentences so that it has a similar meaning to the original sentence.",
    "mq": "Complete each of the following questions.",
    "wb": "Complete each blank with ONE suitable word from the box.",
    "wq": "Write questions for the underlined parts in each of the following sentences.",
    "mt": "Match each item in the left column with its correct definition or meaning in the right column.",
    "tf": "Decide whether the following statements are True (T) or False (F)."
}

def cm_to_pt(cm: float) -> float:
    return float(cm) * 28.346456692913385

def get_instruction_map():
    config_file = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "exercise_config.json")
    if os.path.exists(config_file):
        try:
            with open(config_file, "r", encoding="utf-8") as f:
                custom_map = json.load(f)
                return {**INSTRUCTION_MAP_MCQ, **custom_map}
        except Exception:
            pass
    return INSTRUCTION_MAP_MCQ

def get_instruction_text(ex_type: str, is_mcq: bool = True) -> str:
    config_file = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "exercise_config.json")
    if os.path.exists(config_file):
        try:
            with open(config_file, "r", encoding="utf-8") as f:
                custom_map = json.load(f)
                specific_key = f"{ex_type}_mcq" if is_mcq else f"{ex_type}_non_mcq"
                if specific_key in custom_map:
                    return custom_map[specific_key]
                if ex_type in custom_map:
                    return custom_map[ex_type]
        except Exception:
            pass
    if is_mcq:
        return INSTRUCTION_MAP_MCQ.get(ex_type, "Mark the letter A, B, C, or D on your answer sheet to indicate the answer that best fits each of the following questions.")
    else:
        return INSTRUCTION_MAP_NON_MCQ.get(ex_type, "Answer each of the following questions.")

def format_question_label(q_num: Any, fmt: str = None, default_prefix: str = "") -> str:
    if q_num is None or str(q_num).strip() == "":
        return ""
    q_str = str(q_num).strip()
    if not fmt:
        if default_prefix:
            return f"{default_prefix}{q_str}: "
        return f"{q_str}. "
    
    fmt = fmt.strip()
    if fmt == "1.":
        return f"{q_str}. "
    elif fmt == "(1)":
        return f"({q_str}) "
    elif fmt == "a.":
        return f"{q_str.lower()}. "
    elif fmt == "A.":
        return f"{q_str.upper()}. "
    elif fmt == "a)":
        return f"{q_str.lower()}) "
    else:
        last_char = fmt[-1] if fmt[-1] in '.:)' else ' '
        return f"{q_str}{last_char} "

def parse_text_formatting(text: str) -> List[Dict[str, Any]]:
    import re
    if not text:
        return []
        
    pattern = re.compile(
        r'(?P<err>\[(?P<err_txt>[^\]]+)\]\((?P<err_let>[A-D])\))|'
        r'(?P<bold_under1>\*\*\[(?P<bu_txt1>[^\]]+)\]\*\*)|'
        r'(?P<bold_under2>\[\*\*(?P<bu_txt2>[^*]+)\*\*\])|'
        r'(?P<bold>\*\*(?P<b_txt>[^*]+)\*\*)|'
        r'(?P<italic>\*(?P<i_txt>[^*]+)\*)|'
        r'(?P<under>\[(?P<u_txt>[^\]]+)\])'
    )
    
    segments = []
    last_idx = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            segments.append({
                "text": text[last_idx:start],
                "bold": False,
                "italic": False,
                "underline": False
            })
        
        gd = match.groupdict()
        if gd['err']:
            segments.append({"text": gd['err_txt'], "bold": False, "italic": False, "underline": True})
            segments.append({"text": f" ({gd['err_let']})", "bold": True, "italic": False, "underline": False})
        elif gd['bold_under1'] or gd['bold_under2']:
            txt = gd['bu_txt1'] if gd['bold_under1'] else gd['bu_txt2']
            segments.append({"text": txt, "bold": True, "italic": False, "underline": True})
        elif gd['bold']:
            segments.append({"text": gd['b_txt'], "bold": True, "italic": False, "underline": False})
        elif gd['italic']:
            segments.append({"text": gd['i_txt'], "bold": False, "italic": True, "underline": False})
        elif gd['under']:
            segments.append({"text": gd['u_txt'], "bold": False, "italic": False, "underline": True})
            
        last_idx = end
        
    if last_idx < len(text):
        segments.append({
            "text": text[last_idx:],
            "bold": False,
            "italic": False,
            "underline": False
        })
        
    return segments

def clean_option_text(opt: Any) -> str:
    """Strips leading option prefixes like A., B., C., D., A), B), A:, B:, or repeated A. B. from option text."""
    import re
    if opt is None:
        return ""
    if not isinstance(opt, str):
        opt = str(opt)
    opt = opt.strip()
    return re.sub(r'^(\s*(?:\*\*)?\s*[A-Ea-e]\s*[\.\:\)]\s*(?:\*\*)?\s*)+', '', opt).strip()


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def get_unit_name(grade: str, unit: str) -> str:
    config_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "unit_config.json")
    if not os.path.exists(config_path):
        config_path = os.path.join(os.getcwd(), "unit_config.json")
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                g_key = str(grade).strip()
                u_key = str(unit).strip()
                if g_key in data and u_key in data[g_key]:
                    return data[g_key][u_key]
        except Exception:
            pass
    return ""


class WordDocumentCompilerPyWin32:
    """Real-time MS Word COM Document Compiler using pywin32."""
    def __init__(self, settings: Dict[str, Any] = None):
        self.settings = settings or {}
        self.is_answer_key = False
        self.word = None
        self.doc = None

    def _init_word(self):
        try:
            pythoncom.CoInitialize()
        except Exception:
            pass
            
        self.word = win32com.client.Dispatch("Word.Application")
        self.word.Visible = False
        self.word.DisplayAlerts = 0
        self.doc = self.word.Documents.Add()
        
        self._configure_page()
        self._add_page_numbering()

    def _configure_page(self):
        top_cm = self.settings.get("margin_top", 2.0)
        bottom_cm = self.settings.get("margin_bottom", 2.0)
        left_cm = self.settings.get("margin_left", 3.0)
        right_cm = self.settings.get("margin_right", 1.5)
        
        ps = self.doc.PageSetup
        ps.PageWidth = cm_to_pt(21.0)
        ps.PageHeight = cm_to_pt(29.7)
        ps.TopMargin = cm_to_pt(top_cm)
        ps.BottomMargin = cm_to_pt(bottom_cm)
        ps.LeftMargin = cm_to_pt(left_cm)
        ps.RightMargin = cm_to_pt(right_cm)

    def _add_page_numbering(self):
        font_name = self.settings.get("font_name", "Times New Roman")
        font_size = self.settings.get("font_size", 12.0)
        
        for section in self.doc.Sections:
            footer = section.Footers(1)
            footer.Range.ParagraphFormat.Alignment = 1
            footer.Range.Font.Name = font_name
            footer.Range.Font.Size = font_size
            footer.Range.Text = ""
            
            sel_range = footer.Range
            self.doc.Fields.Add(Range=sel_range, Type=-1, Text="PAGE")
            footer.Range.InsertAfter("/")
            end_range = footer.Range
            end_range.Collapse(0)
            self.doc.Fields.Add(Range=end_range, Type=-1, Text="NUMPAGES")

    def write_run(self, text: str, bold: bool = False, italic: bool = False, underline: bool = False, font_name: str = "Times New Roman", font_size: float = 12.0, color_rgb: int = None, highlight_index: int = None):
        sel = self.word.Selection
        sel.Font.Name = font_name
        sel.Font.Size = font_size
        sel.Font.Bold = bool(bold)
        sel.Font.Italic = bool(italic)
        sel.Font.Underline = 1 if underline else 0
        if color_rgb is not None:
            sel.Font.Color = color_rgb
        if highlight_index is not None:
            sel.Range.HighlightColorIndex = highlight_index
        sel.TypeText(text)
        sel.Font.Bold = 0
        sel.Font.Italic = 0
        sel.Font.Underline = 0
        sel.Font.ColorIndex = 0
        sel.Range.HighlightColorIndex = 0

    def add_instruction_header(self, text: str):
        space_before = self.settings.get("header_space_before", 14.0)
        space_after = self.settings.get("header_space_after", 8.0)
        font_size = self.settings.get("font_size", 12.0)
        
        sel = self.word.Selection
        sel.ParagraphFormat.SpaceBefore = space_before
        sel.ParagraphFormat.SpaceAfter = space_after
        sel.ParagraphFormat.KeepWithNext = True
        self.write_run(text, bold=True, font_size=font_size)
        sel.TypeParagraph()

    def add_question(self, q_num: Any, q_text: str, fmt: str = None, layout: str = None, ind: int = 0, default_prefix: str = ""):
        space_before = self.settings.get("question_space_before", 6.0)
        space_after = self.settings.get("question_space_after", 4.0)
        
        sel = self.word.Selection
        sel.ParagraphFormat.SpaceBefore = space_before
        sel.ParagraphFormat.SpaceAfter = space_after
        sel.ParagraphFormat.KeepWithNext = True
        if ind > 0:
            sel.ParagraphFormat.LeftIndent = cm_to_pt(0.5 * ind)
        else:
            sel.ParagraphFormat.LeftIndent = 0
        
        q_label = format_question_label(q_num, fmt=fmt, default_prefix=default_prefix)
        if q_label:
            self.write_run(q_label, bold=True)
            
        segments = parse_text_formatting(q_text or "")
        for seg in segments:
            self.write_run(seg["text"], bold=seg["bold"], italic=seg["italic"], underline=seg["underline"])
            
        if layout == "blank_right":
            self.write_run("   " + "_" * 25, bold=False)
        elif layout == "tf":
            self.write_run("   ( T  /  F )", bold=True)
            
        sel.TypeParagraph()

    def add_options_grid(self, options: List[str], exercise_type: str, correct_ans: Any = None, cols_override: int = None):
        if not options:
            return
        options = [clean_option_text(opt) for opt in options]
            
        correct_idx = -1
        if correct_ans:
            c_str = str(correct_ans).strip().upper()
            if len(c_str) == 1 and 'A' <= c_str <= 'E':
                correct_idx = ord(c_str) - ord('A')
            elif c_str in ['1', '2', '3', '4', '5']:
                correct_idx = int(c_str) - 1

        if cols_override in [1, 2, 3, 4]:
            cols = cols_override
        else:
            max_len = max(len(str(opt)) for opt in options) if options else 0
            if exercise_type in ["sg", "nt"]:
                cols = 1
            elif exercise_type in ["pr", "st"] or max_len < 15:
                cols = 4
            elif max_len < 35:
                cols = 2
            else:
                cols = 1

        left_indent_cm = self.settings.get("options_left_indent", 0.5)
        space_before = self.settings.get("options_space_before", 0.0)
        space_after = self.settings.get("options_space_after", 3.0)
        
        printable_width_cm = 21.0 - self.settings.get("margin_left", 3.0) - self.settings.get("margin_right", 1.5)
        remaining_width_cm = printable_width_cm - left_indent_cm
        sel = self.word.Selection

        if cols == 1:
            for idx, opt in enumerate(options):
                prefix = chr(65 + idx)
                sel.ParagraphFormat.LeftIndent = cm_to_pt(left_indent_cm)
                sel.ParagraphFormat.SpaceBefore = space_before
                sel.ParagraphFormat.SpaceAfter = space_after
                
                is_curr_ans = (idx == correct_idx and self.is_answer_key)
                color = 255 if is_curr_ans else None
                highlight = 7 if is_curr_ans else None
                
                self.write_run(f"{prefix}. ", bold=True, color_rgb=color, highlight_index=highlight)
                segments = parse_text_formatting(opt)
                for seg in segments:
                    self.write_run(seg["text"], bold=seg["bold"], italic=seg["italic"], underline=seg["underline"], color_rgb=color, highlight_index=highlight)
                sel.TypeParagraph()

        elif cols in [2, 4]:
            col_width = remaining_width_cm / cols
            sel.ParagraphFormat.TabStops.ClearAll()
            for c_i in range(1, cols):
                sel.ParagraphFormat.TabStops.Add(Position=cm_to_pt(left_indent_cm + col_width * c_i), Alignment=0)
            
            num_rows = math.ceil(len(options) / cols)
            for r in range(num_rows):
                sel.ParagraphFormat.LeftIndent = cm_to_pt(left_indent_cm)
                sel.ParagraphFormat.SpaceBefore = space_before
                sel.ParagraphFormat.SpaceAfter = space_after
                
                for c in range(cols):
                    idx = r * cols + c
                    if idx < len(options):
                        prefix = chr(65 + idx)
                        is_curr_ans = (idx == correct_idx and self.is_answer_key)
                        color = 255 if is_curr_ans else None
                        highlight = 7 if is_curr_ans else None
                        
                        self.write_run(f"{prefix}. ", bold=True, color_rgb=color, highlight_index=highlight)
                        for seg in parse_text_formatting(options[idx]):
                            self.write_run(seg["text"], bold=seg["bold"], italic=seg["italic"], underline=seg["underline"], color_rgb=color, highlight_index=highlight)
                    if c < cols - 1:
                        sel.TypeText("\t")
                sel.TypeParagraph()
            sel.ParagraphFormat.TabStops.ClearAll()

    def add_word_box(self, words: List[str]):
        if not words:
            return
        N = len(words)
        cols = 4 if N >= 8 else (3 if N >= 5 else (2 if N >= 3 else N))
        left_margin_cm = self.settings.get("margin_left", 3.0)
        right_margin_cm = self.settings.get("margin_right", 1.5)
        printable_width_cm = 21.0 - left_margin_cm - right_margin_cm
        
        max_word_len = max(len(str(w)) for w in words) if words else 8
        col_width_cm = max(2.8, (max_word_len * 0.18) + 1.0)
        box_width_cm = min(printable_width_cm, cols * col_width_cm)
        
        table = self.doc.Tables.Add(Range=self.word.Selection.Range, NumRows=1, NumColumns=1)
        table.Alignment = 1
        table.Columns(1).Width = cm_to_pt(box_width_cm)
        cell = table.Cell(1, 1)
        cell.Borders.Enable = True
        
        tab_col_w_cm = box_width_cm / cols
        cell.Range.ParagraphFormat.TabStops.ClearAll()
        for i in range(1, cols):
            cell.Range.ParagraphFormat.TabStops.Add(Position=cm_to_pt(tab_col_w_cm * i), Alignment=0)
            
        for i in range(0, N, cols):
            chunk = words[i:i + cols]
            for idx_w, w_text in enumerate(chunk):
                self.write_run(w_text, italic=True)
                if idx_w < len(chunk) - 1:
                    self.word.Selection.TypeText("\t")
            self.word.Selection.TypeParagraph()
        
        self.word.Selection.Start = table.Range.End
        self.word.Selection.Collapse(0)
        self.word.Selection.ParagraphFormat.LeftIndent = 0
        self.word.Selection.ParagraphFormat.TabStops.ClearAll()

    def add_test_header(self, grade: str = "", unit: str = "", version_code: str = "", is_test: bool = False):
        font_name = self.settings.get("font_name", "Times New Roman")
        font_size = self.settings.get("font_size", 12.0)
        
        table = self.doc.Tables.Add(Range=self.word.Selection.Range, NumRows=1, NumColumns=2)
        table.Alignment = 1
        table.Columns(1).Width = cm_to_pt(10.5)
        table.Columns(2).Width = cm_to_pt(6.0)
        
        cell1 = table.Cell(1, 1)
        p1 = cell1.Range.Paragraphs(1)
        p1.SpaceBefore = 0
        p1.SpaceAfter = 2
        run1 = cell1.Range
        run1.Text = "Họ và tên: ........................................................\nLớp: ......................................................................."
        run1.Font.Name = font_name
        run1.Font.Size = font_size
        run1.Font.Italic = True
        
        cell2 = table.Cell(1, 2)
        p2 = cell2.Range.Paragraphs(1)
        p2.Alignment = 2
        
        has_version = bool(version_code and str(version_code).strip())
        if is_test and has_version:
            try:
                v_num = int(str(version_code).strip())
                ver_text = f"ĐỀ {v_num}"
            except ValueError:
                ver_text = f"ĐỀ {version_code}"
            
            nested = cell2.Range.Tables.Add(Range=cell2.Range, NumRows=1, NumColumns=1)
            nested.Alignment = 2
            nested.Columns(1).Width = cm_to_pt(2.5)
            nested.Cell(1, 1).Borders.Enable = True
            n_p = nested.Cell(1, 1).Range.Paragraphs(1)
            n_p.Alignment = 1
            n_r = nested.Cell(1, 1).Range
            n_r.Text = ver_text
            n_r.Font.Name = font_name
            n_r.Font.Size = 12.0
            n_r.Font.Bold = True

        self.word.Selection.Start = table.Range.End
        self.word.Selection.Collapse(0)

        unit_str = str(unit).strip() if unit else ""
        grade_str = str(grade).strip() if grade else ""
        title_text = ""
        if unit_str and unit_str != "0":
            unit_clean = unit_str.upper()
            if unit_clean.startswith("UNIT"):
                title_text = unit_clean
            else:
                unit_name = get_unit_name(grade_str, unit_str)
                title_text = f"UNIT {unit_clean}: {unit_name.upper()}" if unit_name else f"UNIT {unit_clean}"
        elif grade_str and grade_str != "0":
            title_text = f"GRADE {grade_str}"
            
        if title_text:
            p_unit = self.word.Selection
            p_unit.ParagraphFormat.Alignment = 1
            p_unit.ParagraphFormat.SpaceBefore = 18 if (is_test and has_version) else 6
            p_unit.ParagraphFormat.SpaceAfter = 12
            p_unit.ParagraphFormat.KeepWithNext = True
            self.write_run(title_text, bold=True, font_size=12.0)
            p_unit.TypeParagraph()
            p_unit.ParagraphFormat.Alignment = 0

    def collect_answers(self, exercises: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        answers = []
        def add_ans(q_num: Any, ans: Any):
            if not q_num or ans is None or ans == "":
                return
            if isinstance(ans, list):
                ans_str = ", ".join(str(x) for x in ans)
            else:
                ans_str = str(ans).strip()
            if ans_str:
                answers.append({"q": str(q_num).strip(), "a": ans_str})

        for ex in exercises:
            ex_type = ex.get("t")
            if ex.get("k") and isinstance(ex.get("k"), list):
                for sub in ex.get("k"):
                    add_ans(sub.get("q"), sub.get("a"))
            elif ex_type == "mt" and ex.get("a"):
                ans_val = ex.get("a")
                if isinstance(ans_val, list):
                    for idx_item, char_ans in enumerate(ans_val):
                        add_ans(idx_item + 1, char_ans)
                else:
                    add_ans(ex.get("q", "1"), ans_val)
            else:
                add_ans(ex.get("q"), ex.get("a"))
        
        def get_q_num(x):
            try:
                return int(x["q"])
            except ValueError:
                return 999
        answers.sort(key=get_q_num)
        return answers

    def add_answer_key(self, exercises: List[Dict[str, Any]]):
        answers = self.collect_answers(exercises)
        if not answers:
            return
            
        self.word.Selection.InsertBreak(7) # wdPageBreak
        
        sel = self.word.Selection
        sel.ParagraphFormat.SpaceBefore = 12
        sel.ParagraphFormat.SpaceAfter = 12
        sel.ParagraphFormat.KeepWithNext = True
        self.write_run("ANSWER KEY", bold=True, font_size=14.0)
        sel.TypeParagraph()
        
        N = len(answers)
        cols = 5
        rows = math.ceil(N / cols)
        
        table = self.doc.Tables.Add(Range=sel.Range, NumRows=rows, NumColumns=cols)
        table.Alignment = 1
        
        for r in range(rows):
            for c in range(cols):
                idx = c * rows + r
                if idx < N:
                    item = answers[idx]
                    cell = table.Cell(r + 1, c + 1)
                    cell.Range.Text = f"{item['q']}. {item['a']}"
                    cell.Range.Font.Name = self.settings.get("font_name", "Times New Roman")
                    cell.Range.Font.Size = 11.0
                    cell.Range.Font.Bold = True

    def compile_exercises(self, exercises: List[Dict[str, Any]], grade: str = "", unit: str = "", version_code: str = "", include_answer_key: bool = True):
        if isinstance(exercises, dict):
            for key in ["data", "exercises", "questions", "list"]:
                if key in exercises and isinstance(exercises[key], list):
                    exercises = exercises[key]
                    break
            else:
                if "t" in exercises or "q" in exercises:
                    exercises = [exercises]
                else:
                    raise ValueError("JSON dictionary does not contain a list of exercises.")

        if grade or unit or version_code:
            self.add_test_header(grade, unit, version_code)

        current_block_key = None
        for ex in exercises:
            ex_type = ex.get("t", "mq")
            options = ex.get("o", [])
            if not isinstance(options, list):
                options = []
            is_mcq = (len(options) >= 2 and ex_type in ["pr", "st", "sy", "an", "er", "cz", "rd", "sg", "nt", "mq"])
            
            title_prefix = ex.get("title_prefix") or ex.get("prefix") or ex.get("title_num")
            custom_inst = ex.get("instruction") or ex.get("title")
            if not custom_inst and ex.get("x"):
                custom_inst = ex.get("x")
            
            block_key = (ex_type, is_mcq, title_prefix, custom_inst)
            if block_key != current_block_key:
                current_block_key = block_key
                inst_body = custom_inst if custom_inst else get_instruction_text(ex_type, is_mcq)
                if inst_body:
                    final_header = f"{str(title_prefix).strip()} {inst_body}" if (title_prefix and str(title_prefix).strip()) else inst_body
                    self.add_instruction_header(final_header)

            fmt = ex.get("fmt")
            layout = ex.get("layout")
            cols_override = ex.get("cols")
            ind = ex.get("ind", 0)

            if ex.get("w"):
                self.add_word_box(ex.get("w", []))

            b_body = ex.get("b")
            if b_body:
                paragraphs = b_body if isinstance(b_body, list) else [b_body]
                for paragraph_text in paragraphs:
                    sel = self.word.Selection
                    sel.ParagraphFormat.SpaceBefore = 4.0
                    sel.ParagraphFormat.SpaceAfter = 6.0
                    sel.ParagraphFormat.FirstLineIndent = cm_to_pt(0.75)
                    for seg in parse_text_formatting(paragraph_text):
                        self.write_run(seg["text"], bold=seg["bold"], italic=seg["italic"], underline=seg["underline"])
                    sel.TypeParagraph()

            if ex.get("k") and isinstance(ex.get("k"), list):
                for sub in ex.get("k", []):
                    sub_q = sub.get("q")
                    sub_x = sub.get("x", "")
                    sub_o = sub.get("o", [])
                    sub_a = sub.get("a", "")
                    sub_fmt = sub.get("fmt", fmt)
                    sub_layout = sub.get("layout", layout)
                    sub_cols = sub.get("cols", cols_override)
                    sub_ind = sub.get("ind", ind)

                    def_pref = ""
                    self.add_question(sub_q, sub_x, fmt=sub_fmt, layout=sub_layout, ind=sub_ind, default_prefix=def_pref)
                    
                    if sub_o and isinstance(sub_o, list):
                        self.add_options_grid(sub_o, ex_type, correct_ans=sub_a, cols_override=sub_cols)
                    
                    if ex_type == "wq":
                        sel = self.word.Selection
                        sel.ParagraphFormat.LeftIndent = cm_to_pt(0.5)
                        self.write_run("(A:) " + "_" * 50, bold=True)
                        sel.TypeParagraph()

            elif not b_body:
                q_num = ex.get("q")
                q_text = ex.get("x", "")
                options = ex.get("o", [])
                ans = ex.get("a", "")
                
                def_pref = ""

                if ex_type == "ro" and options and not (is_mcq and len(options) == 4):
                    self.add_question(q_num, q_text, fmt=fmt, layout=layout, ind=ind, default_prefix=def_pref)
                    for idx_item, item in enumerate(options):
                        sel = self.word.Selection
                        sel.ParagraphFormat.LeftIndent = cm_to_pt(1.0)
                        for seg in parse_text_formatting(f"{idx_item + 1}. {item}"):
                            self.write_run(seg["text"], bold=seg["bold"], italic=seg["italic"], underline=seg["underline"])
                        sel.TypeParagraph()
                else:
                    self.add_question(q_num, q_text, fmt=fmt, layout=layout, ind=ind, default_prefix=def_pref)
                    if options and isinstance(options, list):
                        self.add_options_grid(options, ex_type, correct_ans=ans, cols_override=cols_override)
                    
                    if ex_type == "wq":
                        sel = self.word.Selection
                        sel.ParagraphFormat.LeftIndent = cm_to_pt(0.5)
                        self.write_run("(A:) " + "_" * 50, bold=True)
                        sel.TypeParagraph()

        if include_answer_key:
            self.add_answer_key(exercises)

    def compile(self, exercises: List[Dict[str, Any]], output_filepath: Any, grade: str = "", unit: str = "", version_code: str = "", include_answer_key: bool = True, is_answer_key: bool = False, is_test: bool = False):
        self.is_answer_key = is_answer_key
        try:
            self._init_word()
            if grade or unit or version_code or is_test:
                self.add_test_header(grade=grade, unit=unit, version_code=version_code, is_test=is_test)
            self.compile_exercises(exercises, grade=grade, unit=unit, version_code=version_code, include_answer_key=include_answer_key)
            abs_path = os.path.abspath(output_filepath)
            os.makedirs(os.path.dirname(abs_path), exist_ok=True)
            self.doc.SaveAs2(abs_path, FileFormat=16)
        finally:
            if self.doc:
                try: self.doc.Close(False)
                except Exception: pass
            if self.word:
                try: self.word.Quit()
                except Exception: pass


class WordDocumentCompilerDocx:
    """Fallback XML document compiler using python-docx."""
    def __init__(self, settings: Dict[str, Any] = None):
        self.settings = settings or {}
        self.doc = Document()
        self.is_answer_key = False
        self._configure_page()
        self._configure_styles()
        self._add_page_numbering()

    def _add_page_numbering(self):
        font_name = self.settings.get("font_name", "Times New Roman")
        font_size = self.settings.get("font_size", 12.0)
        
        for section in self.doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.alignment = 1
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            p.text = ""
            
            def add_field(p_elem, field_name):
                fldSimple = OxmlElement('w:fldSimple')
                fldSimple.set(qn('w:instr'), field_name)
                
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rPr.append(rFonts)
                
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(font_size * 2)))
                rPr.append(sz)
                
                r.append(rPr)
                
                t = OxmlElement('w:t')
                t.text = "1"
                r.append(t)
                
                fldSimple.append(r)
                p_elem._p.append(fldSimple)
                
            add_field(p, 'PAGE')
            r_sep = p.add_run("/")
            r_sep.font.name = font_name
            r_sep.font.size = Pt(font_size)
            add_field(p, 'NUMPAGES')

    def _configure_page(self):
        top_cm = self.settings.get("margin_top", 2.0)
        bottom_cm = self.settings.get("margin_bottom", 2.0)
        left_cm = self.settings.get("margin_left", 3.0)
        right_cm = self.settings.get("margin_right", 1.5)
        
        for section in self.doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(top_cm)
            section.bottom_margin = Cm(bottom_cm)
            section.left_margin = Cm(left_cm)
            section.right_margin = Cm(right_cm)

    def _configure_styles(self):
        font_name = self.settings.get("font_name", "Times New Roman")
        font_size = self.settings.get("font_size", 12.0)
        line_spacing = self.settings.get("line_spacing", 1.15)
        space_after = self.settings.get("space_after", 6.0)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        p_format = style.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(space_after)
        p_format.line_spacing = line_spacing

    def add_instruction_header(self, text: str):
        space_before = self.settings.get("header_space_before", 14.0)
        space_after = self.settings.get("header_space_after", 8.0)
        font_size = self.settings.get("font_size", 12.0)
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)

    def add_question(self, q_num: Any, q_text: str, fmt: str = None, layout: str = None, ind: int = 0, default_prefix: str = ""):
        space_before = self.settings.get("question_space_before", 6.0)
        space_after = self.settings.get("question_space_after", 4.0)
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        if ind > 0:
            p.paragraph_format.left_indent = Cm(0.5 * ind)
            
        q_label = format_question_label(q_num, fmt=fmt, default_prefix=default_prefix)
        if q_label:
            run_q = p.add_run(q_label)
            run_q.bold = True
        
        segments = parse_text_formatting(q_text or "")
        for seg in segments:
            run = p.add_run(seg["text"])
            if seg["bold"]: run.bold = True
            if seg["italic"]: run.italic = True
            if seg["underline"]: run.underline = True

        if layout == "blank_right":
            left_margin_cm = self.settings.get("margin_left", 3.0)
            right_margin_cm = self.settings.get("margin_right", 1.5)
            printable_width_cm = 21.0 - left_margin_cm - right_margin_cm - (0.5 * ind)
            
            p.paragraph_format.tab_stops.add_tab_stop(Cm(printable_width_cm), WD_TAB_ALIGNMENT.RIGHT)
            r_tab = p.add_run("\t" + "_" * 22)
            r_tab.bold = False
        elif layout == "tf":
            run_tf = p.add_run("   ( T  /  F )")
            run_tf.bold = True

        return p

    def add_matching_grid(self, left_items: List[str], right_items: List[str], fmt: str = None):
        if not left_items and not right_items:
            return
        left_margin_cm = self.settings.get("margin_left", 3.0)
        right_margin_cm = self.settings.get("margin_right", 1.5)
        printable_width_cm = 21.0 - left_margin_cm - right_margin_cm
        
        max_left_len = max((len(clean_option_text(str(item))) for item in left_items), default=8)
        left_col_cm = min(5.5, max(3.2, (max_left_len * 0.20) + 1.2))
        right_col_cm = max(8.0, printable_width_cm - left_col_cm)
        
        table = self.doc.add_table(rows=max(len(left_items), len(right_items)), cols=2)
        table.alignment = 1
        table.autofit = False
        
        for r_idx in range(max(len(left_items), len(right_items))):
            if r_idx < len(left_items):
                cell_left = table.cell(r_idx, 0)
                cell_left.width = Cm(left_col_cm)
                p_l = cell_left.paragraphs[0]
                p_l.paragraph_format.space_before = Pt(2)
                p_l.paragraph_format.space_after = Pt(2)
                l_num = format_question_label(r_idx + 1, fmt=fmt or "1.", default_prefix="")
                run_n = p_l.add_run(l_num)
                run_n.bold = True
                for seg in parse_text_formatting(str(left_items[r_idx])):
                    r = p_l.add_run(seg["text"])
                    if seg["bold"]: r.bold = True
                    if seg["italic"]: r.italic = True
                    if seg["underline"]: r.underline = True

            if r_idx < len(right_items):
                cell_right = table.cell(r_idx, 1)
                cell_right.width = Cm(right_col_cm)
                p_r = cell_right.paragraphs[0]
                p_r.paragraph_format.space_before = Pt(2)
                p_r.paragraph_format.space_after = Pt(2)
                for seg in parse_text_formatting(str(right_items[r_idx])):
                    r = p_r.add_run(seg["text"])
                    if seg["bold"]: r.bold = True
                    if seg["italic"]: r.italic = True
                    if seg["underline"]: r.underline = True

        p_space = self.doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(0)
        p_space.paragraph_format.space_after = Pt(4)

    def add_options_grid(self, options: List[str], exercise_type: str, correct_ans: Any = None, cols_override: int = None):
        if not options:
            return
        options = [clean_option_text(opt) for opt in options]
            
        correct_idx = -1
        if correct_ans:
            c_str = str(correct_ans).strip().upper()
            if len(c_str) == 1 and 'A' <= c_str <= 'E':
                correct_idx = ord(c_str) - ord('A')
            elif c_str in ['1', '2', '3', '4', '5']:
                correct_idx = int(c_str) - 1

        if cols_override in [1, 2, 3, 4]:
            cols = cols_override
        else:
            max_len = max(len(str(opt)) for opt in options) if options else 0
            if exercise_type in ["sg", "nt"]:
                cols = 1
            elif exercise_type in ["pr", "st"] or max_len < 15:
                cols = 4
            elif max_len < 35:
                cols = 2
            else:
                cols = 1

        left_indent_cm = self.settings.get("options_left_indent", 0.5)
        space_before = self.settings.get("options_space_before", 0.0)
        space_after = self.settings.get("options_space_after", 3.0)
        
        left_margin_cm = self.settings.get("margin_left", 3.0)
        right_margin_cm = self.settings.get("margin_right", 1.5)
        printable_width_cm = 21.0 - left_margin_cm - right_margin_cm
        remaining_width_cm = printable_width_cm - left_indent_cm

        if cols == 1:
            for idx, opt in enumerate(options):
                prefix = chr(65 + idx)
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(left_indent_cm)
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                
                run_prefix = p.add_run(f"{prefix}. ")
                run_prefix.bold = True
                if idx == correct_idx and self.is_answer_key:
                    run_prefix.font.color.rgb = RGBColor(255, 0, 0)
                    run_prefix.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                segments = parse_text_formatting(opt)
                for seg in segments:
                    run = p.add_run(seg["text"])
                    if seg["bold"]: run.bold = True
                    if seg["italic"]: run.italic = True
                    if seg["underline"]: run.underline = True
                    if idx == correct_idx and self.is_answer_key:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        elif cols == 2:
            col_width = remaining_width_cm / 2
            tab_pos = left_indent_cm + col_width
            num_rows = math.ceil(len(options) / 2)
            
            for row in range(num_rows):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(left_indent_cm)
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_pos), WD_TAB_ALIGNMENT.LEFT)
                
                idx1 = row * 2
                if idx1 < len(options):
                    prefix = chr(65 + idx1)
                    run_prefix = p.add_run(f"{prefix}. ")
                    run_prefix.bold = True
                    if idx1 == correct_idx and self.is_answer_key:
                        run_prefix.font.color.rgb = RGBColor(255, 0, 0)
                        run_prefix.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    for seg in parse_text_formatting(options[idx1]):
                        run = p.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                        if idx1 == correct_idx and self.is_answer_key:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        
                r_tab = p.add_run("\t")
                r_tab.font.underline = False
                
                idx2 = row * 2 + 1
                if idx2 < len(options):
                    prefix = chr(65 + idx2)
                    run_prefix = p.add_run(f"{prefix}. ")
                    run_prefix.bold = True
                    if idx2 == correct_idx and self.is_answer_key:
                        run_prefix.font.color.rgb = RGBColor(255, 0, 0)
                        run_prefix.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    for seg in parse_text_formatting(options[idx2]):
                        run = p.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True
                        if idx2 == correct_idx and self.is_answer_key:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        elif cols == 4:
            col_width = remaining_width_cm / 4
            tab1 = left_indent_cm + col_width
            tab2 = left_indent_cm + (col_width * 2)
            tab3 = left_indent_cm + (col_width * 3)
            num_rows = math.ceil(len(options) / 4)

            for row in range(num_rows):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(left_indent_cm)
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.tab_stops.add_tab_stop(Cm(tab1), WD_TAB_ALIGNMENT.LEFT)
                p.paragraph_format.tab_stops.add_tab_stop(Cm(tab2), WD_TAB_ALIGNMENT.LEFT)
                p.paragraph_format.tab_stops.add_tab_stop(Cm(tab3), WD_TAB_ALIGNMENT.LEFT)
                
                for idx_c in range(4):
                    idx = row * 4 + idx_c
                    if idx < len(options):
                        prefix = chr(65 + idx)
                        run_prefix = p.add_run(f"{prefix}. ")
                        run_prefix.bold = True
                        if idx == correct_idx and self.is_answer_key:
                            run_prefix.font.color.rgb = RGBColor(255, 0, 0)
                            run_prefix.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        for seg in parse_text_formatting(options[idx]):
                            run = p.add_run(seg["text"])
                            if seg["bold"]: run.bold = True
                            if seg["italic"]: run.italic = True
                            if seg["underline"]: run.underline = True
                            if idx == correct_idx and self.is_answer_key:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    if idx_c < 3:
                        r_tab = p.add_run("\t")
                        r_tab.font.underline = False
            
    def add_word_box(self, words: List[str]):
        if not words:
            return
        N = len(words)
        cols = 4 if N >= 8 else (3 if N >= 5 else (2 if N >= 3 else N))
            
        left_margin_cm = self.settings.get("margin_left", 3.0)
        right_margin_cm = self.settings.get("margin_right", 1.5)
        printable_width_cm = 21.0 - left_margin_cm - right_margin_cm
        
        max_word_len = max(len(str(w)) for w in words) if words else 8
        col_width_cm = max(2.8, (max_word_len * 0.18) + 1.0)
        box_width_cm = min(printable_width_cm, cols * col_width_cm)
        
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = 1
        table.autofit = False
        table.columns[0].width = Cm(box_width_cm)
        
        cell = table.cell(0, 0)
        cell.width = Cm(box_width_cm)
        cell.vertical_alignment = 1
        
        border_spec = {"sz": 8, "val": "single", "color": "000000", "space": "0"}
        set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)
        
        tab_col_w_cm = box_width_cm / cols
        for i in range(0, N, cols):
            chunk = words[i:i + cols]
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for c_i in range(1, len(chunk)):
                p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_col_w_cm * c_i), WD_TAB_ALIGNMENT.LEFT)
                
            for idx_w, w_text in enumerate(chunk):
                run = p.add_run(w_text)
                run.italic = True
                if idx_w < len(chunk) - 1:
                    r_tab = p.add_run("\t")
                    r_tab.font.italic = False

    def add_test_header(self, grade: str = "", unit: str = "", version_code: str = "", is_test: bool = False):
        font_name = self.settings.get("font_name", "Times New Roman")
        font_size = self.settings.get("font_size", 12.0)
        
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = 1
        table.autofit = False
        table.columns[0].width = Cm(10.5)
        table.columns[1].width = Cm(6.0)
        
        cell_left = table.cell(0, 0)
        cell_left.width = Cm(10.5)
        p_left = cell_left.paragraphs[0]
        p_left.paragraph_format.space_before = Pt(0)
        p_left.paragraph_format.space_after = Pt(2)
        
        run_name = p_left.add_run("Họ và tên: ........................................................\nLớp: .......................................................................")
        run_name.font.name = font_name
        run_name.font.size = Pt(font_size)
        run_name.italic = True
        
        cell_right = table.cell(0, 1)
        cell_right.width = Cm(6.0)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = 2
        p_right.paragraph_format.space_before = Pt(0)
        p_right.paragraph_format.space_after = Pt(2)
        
        has_version = bool(version_code and str(version_code).strip())
        if is_test and has_version:
            try:
                v_num = int(str(version_code).strip())
                ver_text = f"ĐỀ {v_num}"
            except ValueError:
                ver_text = f"ĐỀ {version_code}"
                
            nested_table = cell_right.add_table(rows=1, cols=1)
            nested_table.alignment = 2
            nested_table.autofit = False
            nested_table.columns[0].width = Cm(2.5)
            
            nested_cell = nested_table.cell(0, 0)
            nested_cell.width = Cm(2.5)
            nested_cell.vertical_alignment = 1
            
            p_nested = nested_cell.paragraphs[0]
            p_nested.alignment = 1
            p_nested.paragraph_format.space_before = Pt(6)
            p_nested.paragraph_format.space_after = Pt(6)
            
            run_ver = p_nested.add_run(ver_text)
            run_ver.bold = True
            run_ver.font.size = Pt(12)
            
            border_spec = {"sz": 12, "val": "single", "color": "000000", "space": "0"}
            set_cell_border(
                nested_cell,
                top=border_spec,
                bottom=border_spec,
                left=border_spec,
                right=border_spec
            )
            
        unit_str = str(unit).strip() if unit else ""
        grade_str = str(grade).strip() if grade else ""
        
        title_text = ""
        if unit_str and unit_str != "0":
            unit_clean = unit_str.upper()
            if unit_clean.startswith("UNIT"):
                title_text = unit_clean
            else:
                unit_name = get_unit_name(grade_str, unit_str)
                if unit_name:
                    title_text = f"UNIT {unit_clean}: {unit_name.upper()}"
                else:
                    title_text = f"UNIT {unit_clean}"
        elif grade_str and grade_str != "0":
            title_text = f"GRADE {grade_str}"
            
        if title_text:
            p_unit = self.doc.add_paragraph()
            p_unit.alignment = 1
            p_unit.paragraph_format.space_before = Pt(18 if (is_test and has_version) else 6)
            p_unit.paragraph_format.space_after = Pt(12)
            p_unit.paragraph_format.keep_with_next = True
            
            run_unit = p_unit.add_run(title_text)
            run_unit.bold = True
            run_unit.font.size = Pt(12)

    def collect_answers(self, exercises: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        answers = []
        def add_ans(q_num: Any, ans: Any):
            if not q_num or ans is None or ans == "":
                return
            if isinstance(ans, list):
                ans_str = ", ".join(str(x) for x in ans)
            else:
                ans_str = str(ans).strip()
            if ans_str:
                answers.append({"q": str(q_num).strip(), "a": ans_str})

        for ex in exercises:
            ex_type = ex.get("t")
            if ex.get("k") and isinstance(ex.get("k"), list):
                for sub in ex.get("k"):
                    add_ans(sub.get("q"), sub.get("a"))
            elif ex_type == "mt" and ex.get("a"):
                ans_val = ex.get("a")
                if isinstance(ans_val, list):
                    for idx_item, char_ans in enumerate(ans_val):
                        add_ans(idx_item + 1, char_ans)
                else:
                    add_ans(ex.get("q", "1"), ans_val)
            else:
                add_ans(ex.get("q"), ex.get("a"))
        
        def get_q_num(x):
            try:
                return int(x["q"])
            except ValueError:
                return 999
        answers.sort(key=get_q_num)
        return answers

    def add_answer_key(self, exercises: List[Dict[str, Any]]):
        answers = self.collect_answers(exercises)
        if not answers:
            return
            
        self.doc.add_page_break()
        
        p_head = self.doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(12)
        p_head.paragraph_format.space_after = Pt(12)
        p_head.paragraph_format.keep_with_next = True
        run_head = p_head.add_run("ANSWER KEY")
        run_head.bold = True
        run_head.font.size = Pt(14)
        
        N = len(answers)
        cols = 5
        rows = math.ceil(N / cols)
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = 1
        
        for r in range(rows):
            for c in range(cols):
                idx = c * rows + r
                if idx < N:
                    item = answers[idx]
                    cell = table.cell(r, c)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(f"{item['q']}. {item['a']}")
                    run.bold = True
                    run.font.size = Pt(11)

    def compile_exercises(self, exercises: List[Dict[str, Any]], grade: str = "", unit: str = "", version_code: str = "", include_answer_key: bool = True):
        if isinstance(exercises, dict):
            for key in ["data", "exercises", "questions", "list"]:
                if key in exercises and isinstance(exercises[key], list):
                    exercises = exercises[key]
                    break
            else:
                if "t" in exercises or "q" in exercises:
                    exercises = [exercises]
                else:
                    raise ValueError("JSON dictionary does not contain a list of exercises.")

        if grade or unit or version_code:
            self.add_test_header(grade, unit, version_code)

        current_block_key = None
        for ex in exercises:
            ex_type = ex.get("t", "mq")
            options = ex.get("o", [])
            if not isinstance(options, list):
                options = []
            is_mcq = (len(options) >= 2 and ex_type in ["pr", "st", "sy", "an", "er", "cz", "rd", "sg", "nt", "mq"])
            
            title_prefix = ex.get("title_prefix") or ex.get("prefix") or ex.get("title_num")
            custom_inst = ex.get("instruction") or ex.get("title")
            if not custom_inst and ex.get("x"):
                custom_inst = ex.get("x")
            
            block_key = (ex_type, is_mcq, title_prefix, custom_inst)
            if block_key != current_block_key:
                current_block_key = block_key
                inst_body = custom_inst if custom_inst else get_instruction_text(ex_type, is_mcq)
                if inst_body:
                    final_header = f"{str(title_prefix).strip()} {inst_body}" if (title_prefix and str(title_prefix).strip()) else inst_body
                    self.add_instruction_header(final_header)

            fmt = ex.get("fmt")
            layout = ex.get("layout")
            cols_override = ex.get("cols")
            ind = ex.get("ind", 0)

            if ex.get("w"):
                self.add_word_box(ex.get("w", []))

            b_body = ex.get("b")
            if b_body:
                space_before = self.settings.get("passage_space_before", 4.0)
                space_after = self.settings.get("passage_space_after", 6.0)
                indent_first = self.settings.get("passage_indent_first", 0.75)
                
                paragraphs = b_body if isinstance(b_body, list) else [b_body]
                for paragraph_text in paragraphs:
                    p_passage = self.doc.add_paragraph()
                    p_passage.paragraph_format.space_before = Pt(space_before)
                    p_passage.paragraph_format.space_after = Pt(space_after)
                    p_passage.paragraph_format.first_line_indent = Cm(indent_first)
                    
                    segments = parse_text_formatting(paragraph_text)
                    for seg in segments:
                        run = p_passage.add_run(seg["text"])
                        if seg["bold"]: run.bold = True
                        if seg["italic"]: run.italic = True
                        if seg["underline"]: run.underline = True

            if layout == "match" or ex_type == "mt":
                left_items = ex.get("o", [])
                right_items = ex.get("p", [])
                if left_items or right_items:
                    self.add_matching_grid(left_items, right_items, fmt=fmt)

            if ex.get("k") and isinstance(ex.get("k"), list):
                for sub in ex.get("k", []):
                    sub_q = sub.get("q")
                    sub_x = sub.get("x", "")
                    sub_o = sub.get("o", [])
                    sub_a = sub.get("a", "")
                    sub_fmt = sub.get("fmt", fmt)
                    sub_layout = sub.get("layout", layout)
                    sub_cols = sub.get("cols", cols_override)
                    sub_ind = sub.get("ind", ind)

                    def_pref = ""
                    self.add_question(sub_q, sub_x, fmt=sub_fmt, layout=sub_layout, ind=sub_ind, default_prefix=def_pref)
                    
                    if sub_o and isinstance(sub_o, list):
                        self.add_options_grid(sub_o, ex_type, correct_ans=sub_a, cols_override=sub_cols)
                    
                    if ex_type == "wq":
                        p_ans = self.doc.add_paragraph()
                        p_ans.paragraph_format.space_before = Pt(2)
                        p_ans.paragraph_format.space_after = Pt(6)
                        p_ans.paragraph_format.left_indent = Cm(0.5)
                        r_ans = p_ans.add_run("(A:) " + "_" * 50)
                        r_ans.bold = True

            elif not b_body and layout != "match" and ex_type != "mt":
                q_num = ex.get("q")
                q_text = ex.get("x", "")
                options = ex.get("o", [])
                ans = ex.get("a", "")
                
                def_pref = ""

                if ex_type == "ro" and options and not (is_mcq and len(options) == 4):
                    self.add_question(q_num, q_text, fmt=fmt, layout=layout, ind=ind, default_prefix=def_pref)
                    for idx_item, item in enumerate(options):
                        p_item = self.doc.add_paragraph()
                        p_item.paragraph_format.left_indent = Cm(1.0)
                        p_item.paragraph_format.space_before = Pt(0)
                        p_item.paragraph_format.space_after = Pt(2)
                        for seg in parse_text_formatting(f"{idx_item + 1}. {item}"):
                            r = p_item.add_run(seg["text"])
                            if seg["bold"]: r.bold = True
                            if seg["italic"]: r.italic = True
                            if seg["underline"]: r.underline = True
                else:
                    self.add_question(q_num, q_text, fmt=fmt, layout=layout, ind=ind, default_prefix=def_pref)
                    if options and isinstance(options, list):
                        self.add_options_grid(options, ex_type, correct_ans=ans, cols_override=cols_override)
                    
                    if ex_type == "wq":
                        p_ans = self.doc.add_paragraph()
                        p_ans.paragraph_format.space_before = Pt(2)
                        p_ans.paragraph_format.space_after = Pt(6)
                        p_ans.paragraph_format.left_indent = Cm(0.5)
                        r_ans = p_ans.add_run("(A:) " + "_" * 50)
                        r_ans.bold = True

        if include_answer_key:
            self.add_answer_key(exercises)

    def compile(self, exercises: List[Dict[str, Any]], output_filepath: Any, grade: str = "", unit: str = "", version_code: str = "", include_answer_key: bool = True, is_answer_key: bool = False, is_test: bool = False):
        self.is_answer_key = is_answer_key
        if grade or unit or version_code or is_test:
            self.add_test_header(grade=grade, unit=unit, version_code=version_code, is_test=is_test)
        self.compile_exercises(exercises, grade=grade, unit=unit, version_code=version_code, include_answer_key=include_answer_key)
        
        abs_path = os.path.abspath(output_filepath)
        os.makedirs(os.path.dirname(abs_path), exist_ok=True)
        self.doc.save(abs_path)


class WordDocumentCompiler:
    """Wrapper compiler that uses win32com if available, otherwise python-docx."""
    def __init__(self, settings: Dict[str, Any] = None):
        self.settings = settings or {}

    def convert_docx_to_pdf_soffice(self, docx_path: str, pdf_path: str) -> bool:
        """Converts DOCX to PDF using LibreOffice/soffice if available."""
        import subprocess
        soffice_paths = [
            r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
            "soffice", "libreoffice"
        ]
        exe = None
        for p in soffice_paths:
            if os.path.exists(p) or p in ["soffice", "libreoffice"]:
                exe = p
                break
        if not exe:
            return False
            
        out_dir = os.path.dirname(os.path.abspath(pdf_path))
        try:
            cmd = [exe, "--headless", "--convert-to", "pdf", os.path.abspath(docx_path), "--outdir", out_dir]
            res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30)
            base_name = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            generated_pdf = os.path.join(out_dir, base_name)
            if os.path.exists(generated_pdf) and generated_pdf != os.path.abspath(pdf_path):
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                os.rename(generated_pdf, pdf_path)
            return os.path.exists(pdf_path)
        except Exception as e:
            print("PDF conversion error:", e)
            return False

    def compile(self, exercises: List[Dict[str, Any]], output_filepath: Any, grade: str = "", unit: str = "", version_code: str = "", include_answer_key: bool = True, is_answer_key: bool = False, is_test: bool = False):
        use_win32 = self.settings.get("use_win32_word", False)
        if use_win32 and win32com_available:
            try:
                compiler_win32 = WordDocumentCompilerPyWin32(self.settings)
                compiler_win32.compile(exercises, output_filepath, grade=grade, unit=unit, version_code=version_code, include_answer_key=include_answer_key, is_answer_key=is_answer_key, is_test=is_test)
                return
            except Exception as e:
                print(f"pywin32 compilation failed, falling back to python-docx: {e}")
        
        compiler_docx = WordDocumentCompilerDocx(self.settings)
        compiler_docx.compile(exercises, output_filepath, grade=grade, unit=unit, version_code=version_code, include_answer_key=include_answer_key, is_answer_key=is_answer_key, is_test=is_test)

    def compile_test_versions(self, exercises: List[Dict[str, Any]], num_versions: int = 1, mix_options: bool = True, grade: str = "", unit: str = "", output_dir: str = None):
        if output_dir and os.path.exists(output_dir):
            out_dir = output_dir
        else:
            try:
                from config.settings import BASE_DIR, get_setting
                files_dir = get_setting("files_dir")
                if files_dir and os.path.exists(files_dir):
                    out_dir = files_dir
                else:
                    out_dir = os.path.join(BASE_DIR, "workspace_files")
            except Exception:
                out_dir = os.path.join(os.getcwd(), "workspace_files")

        os.makedirs(out_dir, exist_ok=True)
        timestamp = int(time.time())
        
        num_v = max(1, num_versions or 1)
        files_created = []
        filenames_created = []
        
        for i in range(num_v):
            v_code = f"{101 + i}" if num_v > 1 else ""
            v_suffix = f"_De_{v_code}" if v_code else ""
            fname = f"DeThi_{grade or 'Test'}_{unit or 'Exam'}_{timestamp}{v_suffix}.docx"
            fpath = os.path.join(out_dir, fname)
            
            ex_copy = copy.deepcopy(exercises)
            self.compile(ex_copy, fpath, grade=grade, unit=unit, version_code=v_code, include_answer_key=True, is_test=True)
            files_created.append(fpath)
            filenames_created.append(fname)
            
        return filenames_created[0], files_created[0], filenames_created
