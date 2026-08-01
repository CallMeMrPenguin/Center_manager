import os
import csv
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_margins_2mm(table):
    """Sets cell margins to 2mm (approx 114 twips/dxa) on all sides."""
    tblPr = table._tbl.tblPr
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
        
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), '114') # 2mm ~ 114 dxa
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def set_all_text_black(doc):
    """Traverses paragraphs, tables, headers, and footers to set font color to black."""
    black = RGBColor(0, 0, 0)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = black
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = black
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                run.font.color.rgb = black
        for p in section.footer.paragraphs:
            for run in p.runs:
                run.font.color.rgb = black

from services.combine_and_format import (
    UNIT_NAMES, 
    apply_vocabulary_formatting, 
    make_section_landscape, 
    set_table_width_to_100, 
    set_table_margins, 
    repeat_table_header, 
    update_docx_fields
)

def read_csv_rows(csv_path: str):
    """Helper to read CSV rows by auto-detecting common delimiters."""
    delimiter = ','
    try:
        with open(csv_path, 'r', encoding='utf-8-sig') as f:
            first_line = f.readline()
            if first_line:
                # Remove quoted parts to avoid false detection of delimiters inside quotes
                cleaned_line = re.sub(r'".*?"', '', first_line)
                delimiters = [',', ';', '\t', '/']
                counts = {d: cleaned_line.count(d) for d in delimiters}
                best_delim = max(counts, key=counts.get)
                if counts[best_delim] > 0:
                    delimiter = best_delim
    except Exception as e:
        print(f"Delimiter auto-detection failed for {csv_path}: {e}")
        
    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        reader = csv.reader(f, delimiter=delimiter)
        return list(reader)

def parse_question_bank_csv(csv_path: str):
    """Parses a Question Bank CSV and returns a list of exercises formatted for the compiler."""
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"Question bank file not found: {csv_path}")
        
    rows = read_csv_rows(csv_path)
    if not rows:
        return []
        
    header = [h.strip() for h in rows[0]]
    header_lower = [h.lower() for h in header]
    
    def find_idx(aliases):
        for alias in aliases:
            a_lower = alias.lower()
            if a_lower in header_lower:
                return header_lower.index(a_lower)
        return -1

    q_idx = find_idx(["questions", "question", "câu hỏi", "cau hoi", "nội dung", "stem"])
    t_idx = find_idx(["question_type", "question type", "questiontype", "loại câu", "loai cau", "dạng câu", "dang cau"])
    o1_idx = find_idx(["option 1", "option1", "phương án 1", "phuong an 1", "tùy chọn 1", "lựa chọn 1"])
    o2_idx = find_idx(["option 2", "option2", "phương án 2", "phuong an 2", "tùy chọn 2", "lựa chọn 2"])
    o3_idx = find_idx(["option 3", "option3", "phương án 3", "phuong an 3", "tùy chọn 3", "lựa chọn 3"])
    o4_idx = find_idx(["option 4", "option4", "phương án 4", "phuong an 4", "tùy chọn 4", "lựa chọn 4"])
    a_idx = find_idx(["answer", "đáp án", "dap an", "đáp án đúng"])

    # Position fallback if standard 11 or 12 column layout
    if (q_idx == -1 or o1_idx == -1 or a_idx == -1) and len(header) >= 11:
        if len(header) >= 12:
            q_idx, t_idx, o1_idx, o2_idx, o3_idx, o4_idx, a_idx = 4, 5, 6, 7, 8, 9, 10
        elif len(header) == 11:
            q_idx, t_idx, o1_idx, o2_idx, o3_idx, o4_idx, a_idx = 3, 4, 5, 6, 7, 8, 9

    if q_idx == -1 or o1_idx == -1 or a_idx == -1:
        missing = []
        if q_idx == -1: missing.append("QUESTIONS")
        if o1_idx == -1: missing.append("OPTION 1")
        if a_idx == -1: missing.append("ANSWER")
        raise ValueError(f"CSV is missing required header fields: {', '.join(missing)}")
        
    no_idx = find_idx(["no.", "no", "stt", "câu số"])
    grade_idx = find_idx(["grade", "khối", "khoi", "lớp", "lop"])
    unit_idx = find_idx(["unit", "bài", "bai", "chủ đề"])
    level_idx = find_idx(["level", "mức độ", "muc do", "độ khó", "do kho"])
    freq_idx = find_idx(["frequency", "tần suất", "tan suat"])
    test_type_idx = find_idx(["test_type", "test type", "dạng đề", "dang de", "loại đề"])

    if no_idx == -1 and len(header) >= 12: no_idx = 0
    if grade_idx == -1 and len(header) >= 12: grade_idx = 1
    if unit_idx == -1 and len(header) >= 12: unit_idx = 2
    if test_type_idx == -1 and len(header) >= 12: test_type_idx = 3
    if level_idx == -1 and len(header) >= 12: level_idx = 11

    exercises = []
    for row_idx, row in enumerate(rows[1:], start=2):
        if not row or not any(x.strip() for x in row):
            continue
        if row[0].strip() == '-':
            continue
            
        try:
            q_num = row[no_idx].strip() if (no_idx != -1 and no_idx < len(row)) else str(row_idx - 1)
            q_text = row[q_idx].strip() if (q_idx != -1 and q_idx < len(row)) else ""
            q_type_raw = row[t_idx].strip() if (t_idx != -1 and t_idx < len(row)) else ""
            
            # Map question type
            q_type_raw_lower = q_type_raw.lower()
            if "fill" in q_type_raw_lower or "blank" in q_type_raw_lower or q_type_raw_lower == "fb":
                t = "fb"
            elif "word" in q_type_raw_lower or "form" in q_type_raw_lower or q_type_raw_lower == "wf":
                t = "wf"
            elif "pronun" in q_type_raw_lower or q_type_raw_lower == "pr":
                t = "pr"
            elif "stress" in q_type_raw_lower or q_type_raw_lower == "st":
                t = "st"
            elif "error" in q_type_raw_lower or q_type_raw_lower == "er":
                t = "er"
            elif "synonym" in q_type_raw_lower or q_type_raw_lower == "sy":
                t = "sy"
            elif "antonym" in q_type_raw_lower or q_type_raw_lower == "an":
                t = "an"
            elif "cloze" in q_type_raw_lower or q_type_raw_lower == "cz":
                t = "cz"
            elif "reading" in q_type_raw_lower or q_type_raw_lower == "rd":
                t = "rd"
            elif "reorder" in q_type_raw_lower or q_type_raw_lower == "ro":
                t = "ro"
            elif "rewrite" in q_type_raw_lower or q_type_raw_lower == "rw":
                t = "rw"
            else:
                t = "mq" # default multiple choice
                
            opts = []
            for idx in [o1_idx, o2_idx, o3_idx, o4_idx]:
                if idx != -1 and idx < len(row):
                    val = row[idx].strip()
                    if val:
                        opts.append(val)
                        
            ans = row[a_idx].strip() if (a_idx != -1 and a_idx < len(row)) else ""
            
            grade_val = row[grade_idx].strip() if (grade_idx != -1 and grade_idx < len(row)) else ""
            unit_val = row[unit_idx].strip() if (unit_idx != -1 and unit_idx < len(row)) else ""
            level_val = row[level_idx].strip() if (level_idx != -1 and level_idx < len(row)) else ""
            freq_val = row[freq_idx].strip() if (freq_idx != -1 and freq_idx < len(row)) else ""
            test_type_val = row[test_type_idx].strip() if (test_type_idx != -1 and test_type_idx < len(row)) else ""
            
            ex = {
                "t": t,
                "q": q_num,
                "x": q_text,
                "o": opts,
                "a": ans,
                "grade": grade_val,
                "unit": unit_val,
                "level": level_val,
                "frequency": freq_val,
                "test_type": test_type_val,
                "meta": {
                    "grade": grade_val,
                    "unit": unit_val
                }
            }
                
            exercises.append(ex)
        except Exception as ex:
            print(f"Error parsing row {row_idx}: {ex}")
            
    return exercises

def parse_vocabulary_csv(csv_path: str):
    """Parses a Vocabulary CSV list and returns a list of vocabulary dict entries."""
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"Vocabulary list file not found: {csv_path}")
        
    rows = read_csv_rows(csv_path)
    if not rows:
        return []
        
    header_row_idx = -1
    for idx, row in enumerate(rows):
        cleaned_row = [x.strip().lower() for x in row]
        if "vocabulary" in cleaned_row or "word" in cleaned_row:
            header_row_idx = idx
            break
            
    if header_row_idx == -1:
        raise ValueError("Could not find vocabulary header row ('Vocabulary' column missing).")
        
    header = [x.strip() for x in rows[header_row_idx]]
    header_lower = [x.lower() for x in header]
    
    if "vocabulary" in header_lower:
        vocab_idx = header_lower.index("vocabulary")
    elif "word" in header_lower:
        vocab_idx = header_lower.index("word")
    else:
        raise ValueError("CSV is missing required vocabulary header field: 'Vocabulary' or 'Word'")
        
    try:
        pos_idx = header_lower.index("pos")
    except ValueError:
        raise ValueError("CSV is missing required vocabulary header field: 'POS'")
        
    try:
        ipa_idx = header_lower.index("ipa")
    except ValueError:
        raise ValueError("CSV is missing required vocabulary header field: 'IPA'")
        
    try:
        if "meaning" in header_lower:
            meaning_idx = header_lower.index("meaning")
        elif "meanings" in header_lower:
            meaning_idx = header_lower.index("meanings")
        else:
            raise ValueError()
    except ValueError:
        raise ValueError("CSV is missing required vocabulary header field: 'Meaning'")
        
    # Helper to find index case-insensitively for optional columns
    def get_optional_index(labels):
        for label in labels:
            if label.lower() in header_lower:
                return header_lower.index(label.lower())
        return -1
        
    no_idx = get_optional_index(["no.", "no"])
    grade_idx = get_optional_index(["grade"])
    unit_idx = get_optional_index(["unit"])
    diff_idx = get_optional_index(["difficulty"])
    root_idx = get_optional_index(["root word", "rootword"])
    
    vocab_entries = []
    for row_idx, row in enumerate(rows[header_row_idx+1:], start=header_row_idx+2):
        if not row or not any(x.strip() for x in row):
            continue
            
        try:
            vocab = row[vocab_idx].strip()
            if not vocab:
                continue
                
            entry = {
                "no": row[no_idx].strip() if (no_idx != -1 and no_idx < len(row)) else "",
                "grade": row[grade_idx].strip() if (grade_idx != -1 and grade_idx < len(row)) else "",
                "unit": row[unit_idx].strip() if (unit_idx != -1 and unit_idx < len(row)) else "",
                "vocabulary": vocab,
                "pos": row[pos_idx].strip() if (pos_idx != -1 and pos_idx < len(row)) else "",
                "ipa": row[ipa_idx].strip() if (ipa_idx != -1 and ipa_idx < len(row)) else "",
                "meaning": row[meaning_idx].strip() if (meaning_idx != -1 and meaning_idx < len(row)) else "",
                "difficulty": row[diff_idx].strip() if (diff_idx != -1 and diff_idx < len(row)) else "",
                "root_word": row[root_idx].strip() if (root_idx != -1 and root_idx < len(row)) else ""
            }
            vocab_entries.append(entry)
        except Exception as e:
            print(f"Error parsing vocabulary row {row_idx}: {e}")
            
    return vocab_entries

def generate_vocab_docx_from_parsed(entries, grade, output_path):
    """Generates a complete, styled vocabulary Word Document for a grade from parsed entries."""
    doc = Document()
    
    # 1. Setup default styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Filter entries for this grade
    grade_entries = [e for e in entries if str(e.get("grade")).strip() == str(grade)]
    if not grade_entries:
        return False
        
    # Group by Unit
    units_dict = {}
    for entry in grade_entries:
        unit = str(entry.get("unit")).strip()
        if not unit.isdigit():
            continue
        unit_num = int(unit)
        if unit_num not in units_dict:
            units_dict[unit_num] = []
        units_dict[unit_num].append(entry)
        
    sorted_units = sorted(list(units_dict.keys()))
    
    # Add TOC and Notes at the beginning
    # TOC Title (Centered)
    toc_title_p = doc.add_paragraph()
    toc_title_p.paragraph_format.page_break_before = False
    toc_title_p.style = 'Normal'
    toc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title_p.add_run("MỤC LỤC")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    # TOC Field Code
    toc_p = doc.add_paragraph()
    run_toc = toc_p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run_toc._r.append(fldChar1)
    run_toc._r.append(instrText)
    run_toc._r.append(fldChar2)
    doc.add_paragraph("Updating Table of Contents...")
    run_toc_end = doc.add_paragraph().add_run()
    run_toc_end._r.append(fldChar3)
    
    # Notes
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(12)
    note_p.paragraph_format.space_after = Pt(6)
    note_run = note_p.add_run(
        "*(Lưu ý: Từ bôi đậm là từ thuộc chủ đề nên ghi; từ in nghiêng là word family cân nhắc ghi; "
        "từ không bôi đậm hay in nghiêng là từ cho lớp Kid .1)*"
    )
    note_run.italic = True
    note_run.font.name = 'Times New Roman'
    note_run.font.size = Pt(11)
    
    # Add section break before unit 1
    new_section = doc.add_section()
    
    # Convert subsequent sections to Landscape (TOC page remains Portrait)
    for s in doc.sections[1:]:
        make_section_landscape(s)
        
    # Build units
    for u_idx, unit_num in enumerate(sorted_units):
        unit_entries = units_dict[unit_num]
        
        # Sort by No.
        def get_no(e):
            no = str(e.get("no")).strip()
            return int(no) if no.isdigit() else 999
        unit_entries.sort(key=get_no)
        
        # Topic name
        from config.unit_config import load_unit_config
        unit_names_dyn = load_unit_config()
        topic_name = unit_names_dyn.get(str(grade), {}).get(str(unit_num), "UNIT TOPIC")
        unit_title = f"UNIT {unit_num}: {topic_name}"
        
        # Centered Unit Title
        p_title = doc.add_paragraph()
        p_title.style = 'Heading 1'
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if u_idx > 0:
            p_title.paragraph_format.page_break_before = True
            
        run_title = p_title.add_run(unit_title)
        run_title.font.name = 'Times New Roman'
        run_title.font.size = Pt(16)
        run_title.bold = True
        
        # Add table (Style Table Grid, exclude Difficulty and Root Word, POS instead of Word Class)
        headers = ["No.", "Vocabulary", "POS", "IPA", "Meaning"]
        table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
        set_table_width_to_100(table)
        set_table_margins_2mm(table)
        repeat_table_header(table)
        
        # Set header row (Centered)
        hdr_row = table.rows[0]
        for c_idx, text in enumerate(headers):
            cell = hdr_row.cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run_h = p.add_run(text)
            run_h.bold = True
            run_h.font.name = 'Times New Roman'
            run_h.font.size = Pt(11)
            
        # Add data rows
        for idx, entry in enumerate(unit_entries, start=1):
            row = table.add_row()
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    
            # 1. No.
            row.cells[0].paragraphs[0].text = str(idx)
            
            # 2. Vocabulary
            p_vocab = row.cells[1].paragraphs[0]
            p_vocab.text = entry.get("vocabulary", "")
            apply_vocabulary_formatting(p_vocab, is_single_para=True)
            
            # 3. POS
            row.cells[2].paragraphs[0].text = entry.get("pos", "")
            
            # 4. IPA
            row.cells[3].paragraphs[0].text = entry.get("ipa", "")
            
            # 5. Meaning
            row.cells[4].paragraphs[0].text = entry.get("meaning", "")
            
        # Format column widths
        col_widths = [Cm(1.5), Cm(6.0), Cm(2.0), Cm(5.0), Cm(10.0)]
        for r in table.rows:
            for c_idx, width in enumerate(col_widths):
                r.cells[c_idx].width = width

    # Setup centered page numbers
    if len(doc.sections) > 1:
        sec2 = doc.sections[1]
        sec2.footer.is_linked_to_previous = False
        footer_p = sec2.footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_num = footer_p.add_run()
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(11)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_num._r.append(fldChar1)
        run_num._r.append(instrText)
        run_num._r.append(fldChar2)
        run_num._r.append(fldChar3)
        
        sectPr_sec2 = sec2._sectPr
        pgNumType = sectPr_sec2.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr_sec2.append(pgNumType)
        pgNumType.set(qn('w:start'), '1')
        
    # Force all text to black before saving
    set_all_text_black(doc)
    doc.save(output_path)
    
    # Run COM field updates
    try:
        update_docx_fields(output_path)
    except Exception:
        pass
    return True
