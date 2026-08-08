import os
import re
import json
import time
import math
from datetime import datetime
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import List, Dict, Any, Optional

from config.settings import get_setting
from services.csv_parser import parse_question_bank_csv
from services.docx_parser import convert_docx_to_json
from database.db_manager import (
    get_connection,
    get_students, create_student, update_student, delete_student,
    get_teachers_cm, create_teacher_cm, update_teacher_cm, delete_teacher_cm,
    get_classes, create_class, update_class, delete_class,
    get_class_students, enroll_student_to_class, unenroll_student_from_class, update_class_student_groups,
    get_class_weekly_schedule, add_class_weekly_slot, delete_class_weekly_slot,
    get_class_sessions, add_class_session, update_class_session, delete_class_session,
    get_courses, create_course, update_course, delete_course,
    get_student_scores, upsert_student_score, delete_student_score,
    get_class_attendance_grades, upsert_class_attendance_grades,
    get_analytics_reports, reset_student_grades, get_class_student_predictions
)
from routers.questions import flatten_docx_to_questions

router = APIRouter()

class StudentPayload(BaseModel):
    full_name: str
    nickname: Optional[str] = ""
    gender: Optional[str] = "Nam"
    grade: Optional[str] = "Lớp 6"
    date_of_birth: Optional[str] = None
    enroll_date: Optional[str] = None
    school: Optional[str] = None
    status: Optional[str] = "Đang học"
    father_name: Optional[str] = None
    father_phone: Optional[str] = None
    mother_name: Optional[str] = None
    mother_phone: Optional[str] = None
    address: Optional[str] = None
    notes: Optional[str] = None

class TeacherCMPayload(BaseModel):
    full_name: str
    role: Optional[str] = "Giáo viên"
    date_of_birth: Optional[str] = None
    phone: Optional[str] = None
    notes: Optional[str] = None

class ClassPayload(BaseModel):
    class_name: str
    teacher_id: Optional[int] = None
    grade: Optional[str] = "Lớp 6"
    subject: Optional[str] = None
    room: Optional[str] = None
    status: Optional[str] = "Đang hoạt động"
    color: Optional[str] = "#7c3aed"
    notes: Optional[str] = None

class EnrollPayload(BaseModel):
    student_id: int
    seat_color: Optional[str] = None
    grade_group: Optional[str] = None

class WeeklySlotPayload(BaseModel):
    day_of_week: str
    start_time: str
    duration: int
    notes: Optional[str] = None

class ClassSessionPayload(BaseModel):
    class_id: Optional[int] = None
    date: str
    start_time: str
    duration: int
    status: Optional[str] = "Sắp diễn ra"
    teacher_id: Optional[int] = None
    notes: Optional[str] = ""
    color: Optional[str] = None

class CoursePayload(BaseModel):
    course_name: str
    description: Optional[str] = None
    price: Optional[float] = 0
    duration_weeks: Optional[int] = None
    status: Optional[str] = "Đang mở"

class ScorePayload(BaseModel):
    student_id: int
    class_id: int
    test_title: Optional[str] = None
    test_date: Optional[str] = None
    score_type: str
    score: float
    max_score: Optional[float] = 10
    notes: Optional[str] = None

@router.get("/api/students")
def api_get_students(search: str = "", status: str = ""):
    return get_students(search, status)

@router.post("/api/students")
def api_create_student(payload: StudentPayload):
    try:
        sid = create_student(payload.dict())
        return {"id": sid, "status": "success"}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.put("/api/students/{student_id}")
def api_update_student(student_id: int, payload: StudentPayload):
    update_student(student_id, payload.dict())
    return {"status": "success"}

@router.delete("/api/students/{student_id}")
def api_delete_student(student_id: int):
    delete_student(student_id)
    return {"status": "success"}

@router.get("/api/teachers_cm")
def api_get_teachers_cm(search: str = "", role: str = ""):
    return get_teachers_cm(search, role)

@router.post("/api/teachers_cm")
def api_create_teacher_cm(payload: TeacherCMPayload):
    tid = create_teacher_cm(payload.dict())
    return {"id": tid, "status": "success"}

@router.put("/api/teachers_cm/{teacher_id}")
def api_update_teacher_cm(teacher_id: int, payload: TeacherCMPayload):
    update_teacher_cm(teacher_id, payload.dict())
    return {"status": "success"}

@router.delete("/api/teachers_cm/{teacher_id}")
def api_delete_teacher_cm(teacher_id: int):
    delete_teacher_cm(teacher_id)
    return {"status": "success"}

@router.get("/api/classes")
def api_get_classes(search: str = ""):
    return get_classes(search)

@router.post("/api/classes")
def api_create_class(payload: ClassPayload):
    cid = create_class(payload.dict())
    return {"id": cid, "status": "success"}

@router.put("/api/classes/{class_id}")
def api_update_class(class_id: int, payload: ClassPayload):
    update_class(class_id, payload.dict())
    return {"status": "success"}

@router.delete("/api/classes/{class_id}")
def api_delete_class(class_id: int):
    delete_class(class_id)
    return {"status": "success"}

@router.get("/api/classes/{class_id}/students")
def api_get_class_students(class_id: int):
    return get_class_students(class_id)

@router.post("/api/classes/{class_id}/students")
def api_enroll_student(class_id: int, payload: EnrollPayload):
    enroll_student_to_class(class_id, payload.student_id, payload.seat_color, payload.grade_group)
    return {"status": "success"}

@router.delete("/api/classes/{class_id}/students/{student_id}")
def api_unenroll_student(class_id: int, student_id: int):
    unenroll_student_from_class(class_id, student_id)
    return {"status": "success"}

@router.put("/api/classes/{class_id}/students/{student_id}/groups")
def api_update_student_groups(class_id: int, student_id: int, payload: EnrollPayload):
    update_class_student_groups(class_id, student_id, payload.seat_color, payload.grade_group)
    return {"status": "success"}

@router.get("/api/classes/{class_id}/schedule/weekly")
def api_get_weekly(class_id: int):
    return get_class_weekly_schedule(class_id)

@router.post("/api/classes/{class_id}/schedule/weekly")
def api_add_weekly(class_id: int, payload: WeeklySlotPayload):
    sid = add_class_weekly_slot(class_id, payload.day_of_week, payload.start_time, payload.duration, payload.notes or "")
    return {"id": sid, "status": "success"}

@router.post("/api/classes/{class_id}/schedule/weekly/replace")
def api_replace_weekly_slots(class_id: int, payload: List[WeeklySlotPayload]):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM class_schedule_weekly WHERE class_id = ?", (class_id,))
    for slot in payload:
        cursor.execute("""
            INSERT INTO class_schedule_weekly (class_id, day_of_week, start_time, duration, notes)
            VALUES (?, ?, ?, ?, ?)
        """, (class_id, slot.day_of_week, slot.start_time, slot.duration, slot.notes or ""))
    conn.commit()
    conn.close()
    return {"status": "success"}

@router.delete("/api/classes/{class_id}/schedule/weekly/{slot_id}")
def api_delete_weekly(slot_id: int):
    delete_class_weekly_slot(slot_id)
    return {"status": "success"}

@router.get("/api/classes/{class_id}/schedule/sessions")
def api_get_sessions(class_id: int, month: str = ""):
    return get_class_sessions(class_id, month)

@router.post("/api/classes/{class_id}/schedule/sessions")
def api_add_session(class_id: int, payload: ClassSessionPayload):
    target_cid = class_id if class_id > 0 else (payload.class_id or 1)
    sid = add_class_session(target_cid, payload.date, payload.start_time, payload.duration, payload.status or "Sắp diễn ra", payload.teacher_id, payload.notes or "", payload.color)
    return {"id": sid, "status": "success"}

@router.put("/api/classes/{class_id}/schedule/sessions/{session_id}")
def api_update_session(class_id: int, session_id: int, payload: ClassSessionPayload):
    update_class_session(session_id, payload.dict(), class_id=class_id)
    return {"status": "success"}

@router.delete("/api/classes/{class_id}/schedule/sessions/{session_id}")
def api_delete_session(session_id: int):
    delete_class_session(session_id)
    return {"status": "success"}

@router.get("/api/classes/{class_id}/attendance")
def api_get_attendance(class_id: int, date: str):
    students = get_class_students(class_id)
    existing = get_class_attendance_grades(class_id, date)
    existing_map = {r["student_id"]: r for r in existing}
    student_preds = get_class_student_predictions(class_id)
    
    result = []
    for st in students:
        preds = student_preds.get(st["id"], {
            "pred_c1": 8.5,
            "pred_c2": 8.0,
            "pred_hw": 9.0,
            "predicted_next": 8.5
        })
        if st["id"] in existing_map:
            rec = dict(existing_map[st["id"]])
            rec["student_name"] = st["full_name"]
            rec["pred_c1"] = preds["pred_c1"]
            rec["pred_c2"] = preds["pred_c2"]
            rec["pred_hw"] = preds["pred_hw"]
            result.append(rec)
        else:
            result.append({
                "class_id": class_id,
                "student_id": st["id"],
                "student_name": st["full_name"],
                "date": date,
                "status": "Có mặt",
                "check_1": None,
                "check_2": None,
                "homework": None,
                "notes": "",
                "pred_c1": preds["pred_c1"],
                "pred_c2": preds["pred_c2"],
                "pred_hw": preds["pred_hw"]
            })
    return {"date": date, "records": result}

@router.post("/api/classes/{class_id}/attendance")
def api_save_attendance(class_id: int, payload: Dict[str, Any]):
    date_str = payload.get("date")
    records = payload.get("records", [])
    if not date_str:
        raise HTTPException(status_code=400, detail="Thiếu ngày điểm danh")
    upsert_class_attendance_grades(class_id, date_str, records)
    return {"status": "success"}

@router.post("/api/classes/{class_id}/export/excel")
def api_export_class_excel(class_id: int, payload: Dict[str, Any]):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.worksheet.table import Table, TableStyleInfo
    
    date_str = payload.get("date") or datetime.now().strftime("%Y-%m-%d")
    cls_list = get_classes()
    cls_info = next((c for c in cls_list if c["id"] == class_id), None)
    class_name = cls_info["class_name"] if cls_info else f"Class_{class_id}"
    
    attendance = payload.get("records")
    if not attendance:
        attendance = api_get_attendance(class_id, date_str).get("records", [])
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Báo Cáo Lớp Học"
    
    ws.merge_cells("A1:I1")
    ws["A1"] = f"BÁO CÁO ĐIỂM DANH & ĐIỂM BÀI HỌC - {class_name.upper()} ({date_str})"
    ws["A1"].font = Font(size=14, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill(start_color="1E1B4B", end_color="1E1B4B", fill_type="solid")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36
    
    headers = ["STT", "Họ và Tên", "Điểm Danh", "Check 1", "Check 2", "BTVN", "BTVN - Check 2", "Check 2 - Check 1", "Cần Cố Gắng (Dưới TB)"]
    ws.cell(row=2, column=1, value="")
    ws.row_dimensions[3].height = 26
    for col_idx, h in enumerate(headers, 1):
        ws.cell(row=3, column=col_idx, value=h)
    
    header_fill = PatternFill(start_color="312E81", end_color="312E81", fill_type="solid")
    header_font = Font(name="Times New Roman", color="FFFFFF", bold=True, size=13)
    data_font = Font(name="Times New Roman", size=13)
    thin_border = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    for col_num in range(1, 10):
        cell = ws.cell(row=3, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
        
    def clean_num(val):
        if val is None:
            return 0.0
        val_str = str(val).strip()
        if not val_str or "không" in val_str.lower():
            return 0.0
        match = re.search(r"[-+]?\d*\.\d+|\d+", val_str)
        if match:
            try:
                return float(match.group(0))
            except ValueError:
                return 0.0
        return 0.0

    start_row = 4
    end_row = start_row + len(attendance) - 1 if len(attendance) > 0 else start_row
    avg_row_idx = end_row + 1

    for idx, r in enumerate(attendance, 1):
        curr_row = start_row + idx - 1
        st_name = str(r.get("student_name", ""))
        status_val = str(r.get("status", "Có mặt"))
        c1 = clean_num(r.get("check_1"))
        c2 = clean_num(r.get("check_2"))
        hw = clean_num(r.get("homework"))

        ws.cell(row=curr_row, column=1, value=f"=ROW()-3")
        ws.cell(row=curr_row, column=2, value=st_name)
        ws.cell(row=curr_row, column=3, value=status_val)
        
        c1_cell = ws.cell(row=curr_row, column=4, value=c1)
        c2_cell = ws.cell(row=curr_row, column=5, value=c2)
        hw_cell = ws.cell(row=curr_row, column=6, value=hw)
        
        c1_cell.number_format = '0.0'
        c2_cell.number_format = '0.0'
        hw_cell.number_format = '0.0'

        c7_cell = ws.cell(row=curr_row, column=7, value=f"=ROUNDUP(ABS(F{curr_row}-E{curr_row}), 1)")
        c8_cell = ws.cell(row=curr_row, column=8, value=f"=ROUNDUP(ABS(E{curr_row}-D{curr_row}), 1)")
        c7_cell.number_format = '0.0'
        c8_cell.number_format = '0.0'

        col_9_cell = ws.cell(
            row=curr_row,
            column=9,
            value=f'=IF(C{curr_row}="Vắng mặt", "Vắng mặt", IF(_xlfn.TEXTJOIN(", ", TRUE, IF(AND(D{curr_row}>0, D{curr_row}<D${avg_row_idx}), "Check 1", ""), IF(AND(E{curr_row}>0, E{curr_row}<E${avg_row_idx}), "Check 2", ""), IF(AND(F{curr_row}>0, F{curr_row}<F${avg_row_idx}), "BTVN", ""))="", "Đạt yêu cầu", "⚠️ Cần cố gắng (" & _xlfn.TEXTJOIN(", ", TRUE, IF(AND(D{curr_row}>0, D{curr_row}<D${avg_row_idx}), "Check 1", ""), IF(AND(E{curr_row}>0, E{curr_row}<E${avg_row_idx}), "Check 2", ""), IF(AND(F{curr_row}>0, F{curr_row}<F${avg_row_idx}), "BTVN", "")) & ")"))'
        )

        for col_num in range(1, 10):
            c_cell = ws.cell(row=curr_row, column=col_num)
            c_cell.font = data_font
            c_cell.border = thin_border
            c_cell.alignment = Alignment(horizontal="center", vertical="center")

    # Add Official Excel Table Object (Format as Table - Non-aqua style TableStyleMedium9)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if len(attendance) > 0:
        table_ref = f"A3:I{end_row}"
        tab = Table(displayName=f"ClassTable_{ts}", ref=table_ref)
        tab.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False
        )
        ws.add_table(tab)

        # Colored text for Column I without background fills so Column I shares the exact table row background style
        from openpyxl.formatting.rule import FormulaRule
        red_font = Font(color="B91C1C", bold=True)
        green_font = Font(color="15803D", bold=True)
        grey_font = Font(color="64748B", bold=True)

        rule_red = FormulaRule(
            formula=['NOT(ISERROR(SEARCH("Cần cố gắng", I4)))'],
            font=red_font
        )
        rule_green = FormulaRule(
            formula=['NOT(ISERROR(SEARCH("Đạt yêu cầu", I4)))'],
            font=green_font
        )
        rule_grey = FormulaRule(
            formula=['NOT(ISERROR(SEARCH("Vắng mặt", I4)))'],
            font=grey_font
        )
        ws.conditional_formatting.add(f"I4:I{end_row}", rule_red)
        ws.conditional_formatting.add(f"I4:I{end_row}", rule_green)
        ws.conditional_formatting.add(f"I4:I{end_row}", rule_grey)

    # Average row
    ws.cell(row=avg_row_idx, column=1, value="")
    ws.cell(row=avg_row_idx, column=2, value="Điểm trung bình (Average)")
    ws.cell(row=avg_row_idx, column=3, value="")

    def trunc_1_dec(v):
        return math.floor(v * 10) / 10.0

    if len(attendance) > 0:
        c1_vals = [clean_num(r.get("check_1")) for r in attendance if clean_num(r.get("check_1")) > 0 and str(r.get("status")) != "Vắng mặt"]
        c2_vals = [clean_num(r.get("check_2")) for r in attendance if clean_num(r.get("check_2")) > 0 and str(r.get("status")) != "Vắng mặt"]
        hw_vals = [clean_num(r.get("homework")) for r in attendance if clean_num(r.get("homework")) > 0 and str(r.get("status")) != "Vắng mặt"]
        
        avg_1 = trunc_1_dec(sum(c1_vals) / len(c1_vals)) if c1_vals else 0.0
        avg_2 = trunc_1_dec(sum(c2_vals) / len(c2_vals)) if c2_vals else 0.0
        avg_hw = trunc_1_dec(sum(hw_vals) / len(hw_vals)) if hw_vals else 0.0

        c1_avg_cell = ws.cell(row=avg_row_idx, column=4, value=avg_1)
        c2_avg_cell = ws.cell(row=avg_row_idx, column=5, value=avg_2)
        hw_avg_cell = ws.cell(row=avg_row_idx, column=6, value=avg_hw)

        c1_avg_cell.number_format = '0.0'
        c2_avg_cell.number_format = '0.0'
        hw_avg_cell.number_format = '0.0'

        diff_hw_c2 = trunc_1_dec(abs(avg_hw - avg_2))
        diff_c2_c1 = trunc_1_dec(abs(avg_2 - avg_1))

        c7_avg_cell = ws.cell(row=avg_row_idx, column=7, value=diff_hw_c2)
        c8_avg_cell = ws.cell(row=avg_row_idx, column=8, value=diff_c2_c1)

        c7_avg_cell.number_format = '0.0'
        c8_avg_cell.number_format = '0.0'

        ws.cell(row=avg_row_idx, column=9, value=f"=IF(D{avg_row_idx}>0, \"Đã tính TB lớp\", \"Chưa đủ điểm\")")

    avg_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    avg_font = Font(name="Times New Roman", bold=True, size=13, color="92400E")
    for col_num in range(1, 10):
        c_cell = ws.cell(row=avg_row_idx, column=col_num)
        c_cell.font = avg_font
        c_cell.fill = avg_fill
        c_cell.border = thin_border
        c_cell.alignment = Alignment(horizontal="center", vertical="center")

    summary_labels = [
        ("Check 1", 4, "D"),
        ("Check 2", 5, "E"),
        ("Bài tập về nhà", 6, "F")
    ]

    for idx, (m_label, col_num, col_let) in enumerate(summary_labels):
        r_idx = avg_row_idx + 2 + idx
        ws.cell(row=r_idx, column=1, value="")
        sum_title = ws.cell(
            row=r_idx,
            column=2,
            value=(
                f'="Check 1 dưới TB (< " & TEXT({col_let}{avg_row_idx}, "0.0") & ")"' if m_label == "Check 1" else
                f'="Check 2 dưới TB (< " & TEXT({col_let}{avg_row_idx}, "0.0") & ")"' if m_label == "Check 2" else
                f'="BTVN dưới TB (< " & TEXT({col_let}{avg_row_idx}, "0.0") & ")"'
            )
        )
        sum_title.font = Font(name="Times New Roman", bold=True, color="7F1D1D", size=13)
        sum_title.alignment = Alignment(horizontal="left", vertical="center")
        ws.merge_cells(f"C{r_idx}:I{r_idx}")
        val_cell = ws.cell(
            row=r_idx,
            column=3,
            value=f'=_xlfn.TEXTJOIN(", ", TRUE, _xlfn.FILTER(B{start_row}:B{end_row}, ({col_let}{start_row}:{col_let}{end_row}>0)*({col_let}{start_row}:{col_let}{end_row}<{col_let}{avg_row_idx})*(C{start_row}:C{end_row}<>"Vắng mặt"), "Không có (Tất cả đạt)"))' if len(attendance) > 0 else "Không có (Tất cả đạt)"
        )
        val_cell.font = Font(name="Times New Roman", bold=True, color="1E1E2F", size=13)
        val_cell.alignment = Alignment(horizontal="left", vertical="center")

    # Auto-fit column widths from row 3 (headings) down to bottom summary rows
    total_max_row = avg_row_idx + 5
    for col_idx in range(1, 10):
        col_let = openpyxl.utils.get_column_letter(col_idx)
        max_len = 0
        for r_idx in range(3, total_max_row + 1):
            if r_idx > end_row and col_idx == 3:
                continue
            cell_val = ws.cell(row=r_idx, column=col_idx).value
            val_str = str(cell_val) if cell_val is not None else ""
            if val_str.startswith("="):
                if col_idx == 9:
                    val_str = "⚠️ Cần cố gắng (Check 1, Check 2, BTVN)"
                elif col_idx == 2:
                    val_str = "Check 1 dưới TB (< 10.0)"
                elif col_idx == 1:
                    val_str = "999"
                elif col_idx in (7, 8):
                    val_str = "10.0"
                else:
                    val_str = ""
            if len(val_str) > max_len:
                max_len = len(val_str)

        extra_padding = 12 if col_idx == 9 else (8 if col_idx == 2 else 5)
        min_w = 54 if col_idx == 9 else (36 if col_idx == 2 else (16 if col_idx == 3 else 14))
        ws.column_dimensions[col_let].width = max(max_len + extra_padding, min_w)

    files_dir = get_setting("files_dir")
    os.makedirs(files_dir, exist_ok=True)
    filename = f"ClassReport_{class_name}_{date_str}_{ts}.xlsx"
    filepath = os.path.join(files_dir, filename)
    wb.save(filepath)
    return {"status": "success", "filename": filename, "filepath": filepath}

@router.post("/api/classes/{class_id}/export/docx")
def api_export_class_docx(class_id: int, payload: Dict[str, Any]):
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    date_str = payload.get("date") or datetime.now().strftime("%Y-%m-%d")
    cls_list = get_classes()
    cls_info = next((c for c in cls_list if c["id"] == class_id), None)
    class_name = cls_info["class_name"] if cls_info else f"Class_{class_id}"
    
    attendance = payload.get("records")
    if not attendance:
        attendance = api_get_attendance(class_id, date_str).get("records", [])
    
    doc = docx.Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"BÁO CÁO NGHỈ HỌC & ĐIỂM BÀI HỌC\nLỚP: {class_name.upper()} - NGÀY: {date_str}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 27, 75)
    
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["STT", "Họ và Tên", "Điểm Danh", "Check 1", "Check 2", "BTVN"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    for idx, r in enumerate(attendance, 1):
        c1 = r.get("check_1", 0)
        c2 = r.get("check_2", 0)
        hw = r.get("homework", 0)
        hw_str = "Không BTVN" if hw == 0 else str(hw)
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = r.get("student_name", "")
        row_cells[2].text = r.get("status", "")
        row_cells[3].text = str(c1)
        row_cells[4].text = str(c2)
        row_cells[5].text = hw_str
        
    files_dir = get_setting("files_dir")
    os.makedirs(files_dir, exist_ok=True)
    ts = datetime.now().strftime("%H%M%S")
    filename = f"ClassReport_{class_name}_{date_str}_{ts}.docx"
    filepath = os.path.join(files_dir, filename)
    doc.save(filepath)
    return {"status": "success", "filename": filename, "filepath": filepath}

@router.get("/api/courses")
def api_get_courses(search: str = "", status: str = ""):
    return get_courses(search, status)

@router.post("/api/courses")
def api_create_course(payload: CoursePayload):
    cid = create_course(payload.dict())
    return {"id": cid, "status": "success"}

@router.put("/api/courses/{course_id}")
def api_update_course(course_id: int, payload: CoursePayload):
    update_course(course_id, payload.dict())
    return {"status": "success"}

@router.delete("/api/courses/{course_id}")
def api_delete_course(course_id: int):
    delete_course(course_id)
    return {"status": "success"}

@router.get("/api/scores")
def api_get_scores(class_id: Optional[int] = None, student_id: Optional[int] = None):
    return get_student_scores(class_id, student_id)

@router.post("/api/scores")
def api_upsert_score(payload: ScorePayload):
    sid = upsert_student_score(payload.dict())
    return {"id": sid, "status": "success"}

@router.delete("/api/scores/{score_id}")
def api_delete_score(score_id: int):
    delete_student_score(score_id)
    return {"status": "success"}

def _clean_opt_prefix(val: Any) -> str:
    if not isinstance(val, str):
        return str(val or "")
    s = val.strip()
    cleaned = re.sub(r'^[A-Da-d0-9][.\):\-]\s*', '', s).strip()
    return cleaned if cleaned else s

@router.post("/api/kiemtra/parse")
async def api_parse_kiemtra(file: UploadFile = File(None), raw_json: Optional[str] = Form(None)):
    files_dir = get_setting("files_dir")
    if file:
        filename = file.filename.lower()
        content = await file.read()
        if filename.endswith(".json"):
            try:
                data = json.loads(content.decode("utf-8"))
                if isinstance(data, list):
                    data = {"title": file.filename, "questions": data}
                elif isinstance(data, dict) and "questions" not in data:
                    raw_qs = data.get("exercises", []) or data.get("data", [])
                    data = {"title": file.filename, "questions": raw_qs}
                
                top_inst = data.get("instruction") or data.get("guide") or ""
                normalized_qs = []
                for idx, q in enumerate(data.get("questions", []), 1):
                    raw_x = q.get("x") or q.get("question")
                    if isinstance(raw_x, list):
                        q_text = "\n".join([str(item) for item in raw_x if item])
                    elif isinstance(raw_x, str):
                        q_text = raw_x
                    elif q.get("sentence") and q.get("prompt"):
                        prompt_str = str(q["prompt"]).strip()
                        if not prompt_str.startswith("("):
                            prompt_str = f"({prompt_str})"
                        q_text = f"{str(q['sentence']).strip()}\n{prompt_str}"
                    else:
                        q_text = q.get("sentence") or q.get("passage") or q.get("content") or f"Câu {idx}"

                    opts = q.get("options") or q.get("o") or []
                    cleaned_opts = [_clean_opt_prefix(o) for o in opts] if isinstance(opts, list) else []
                    ans = q.get("answer") or q.get("a") or ""
                    q_type = q.get("type") or ("mcq" if cleaned_opts else "fill")
                    inst = q.get("instruction") or top_inst
                    normalized_qs.append({
                        "id": q.get("id") or q.get("number") or idx,
                        "question": q_text,
                        "instruction": inst,
                        "type": q_type,
                        "options": cleaned_opts,
                        "answer": str(ans),
                        "explanation": q.get("explanation", "")
                    })
                data["questions"] = normalized_qs
                return data
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Invalid JSON file: {e}")
        elif filename.endswith(".docx"):
            temp_path = os.path.join(files_dir, f"temp_kt_{int(time.time())}_{file.filename}")
            with open(temp_path, "wb") as f:
                f.write(content)
            try:
                parsed_ex = convert_docx_to_json(temp_path)
                flat_qs = flatten_docx_to_questions(parsed_ex)
                questions = []
                for idx, q in enumerate(flat_qs, 1):
                    opts = q.get("o", [])
                    cleaned_opts = [_clean_opt_prefix(o) for o in opts] if isinstance(opts, list) else []
                    q_raw_text = q.get("x", "")
                    instruction = q.get("instruction", "")
                    if not instruction and q_raw_text.startswith("[Yêu cầu: "):
                        parts = q_raw_text.split("]\n", 1)
                        if len(parts) == 2:
                            instruction = parts[0].replace("[Yêu cầu: ", "").strip()
                            q_raw_text = parts[1]
                    questions.append({
                        "id": idx,
                        "question": q_raw_text,
                        "instruction": instruction,
                        "type": "mcq" if cleaned_opts else "fill",
                        "options": cleaned_opts,
                        "answer": q.get("a", ""),
                        "explanation": ""
                    })
                return {"title": file.filename, "questions": questions}
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        elif filename.endswith(".csv"):
            temp_path = os.path.join(files_dir, f"temp_kt_{int(time.time())}_{file.filename}")
            with open(temp_path, "wb") as f:
                f.write(content)
            try:
                parsed_q = parse_question_bank_csv(temp_path)
                questions = []
                for idx, q in enumerate(parsed_q, 1):
                    opts = q.get("o", [])
                    cleaned_opts = [_clean_opt_prefix(o) for o in opts] if isinstance(opts, list) else []
                    questions.append({
                        "id": idx,
                        "question": q.get("x", ""),
                        "type": "mcq" if cleaned_opts else "fill",
                        "options": cleaned_opts,
                        "answer": q.get("a", ""),
                        "explanation": ""
                    })
                return {"title": file.filename, "questions": questions}
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        else:
            raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx, .json hoặc .csv")
    elif raw_json:
        try:
            return json.loads(raw_json)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid raw JSON: {e}")
    else:
        raise HTTPException(status_code=400, detail="No file or JSON content provided")

@router.get("/api/reports/grade-analytics")
def api_get_grade_analytics(class_id: Optional[int] = None, student_id: Optional[int] = None):
    return get_analytics_reports(class_id, student_id)

class ResetGradesPayload(BaseModel):
    class_id: Optional[int] = None
    student_id: Optional[int] = None
    from_date: Optional[str] = None
    to_date: Optional[str] = None

@router.post("/api/reports/reset-grades")
def api_reset_grades(payload: ResetGradesPayload):
    count = reset_student_grades(payload.class_id, payload.student_id, payload.from_date, payload.to_date)
    return {"status": "success", "reset_count": count}
