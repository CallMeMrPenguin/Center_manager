import os
import re
import math
from datetime import datetime
from typing import Dict, Any, List, Optional
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.formatting.rule import FormulaRule
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config.settings import get_setting
from database.db_manager import get_classes, get_class_attendance_grades
from database.utils import trunc_1_dec

def clean_num(val: Any) -> float:
    if val is None:
        return 0.0
    val_str = str(val).strip()
    if not val_str or "không" in val_str.lower():
        return 0.0
    match = re.search(r"[-+]?\d*\.\d+|\d+", val_str)
    if match:
        try:
            return float(match.group(0))
        except ValueError:
            return 0.0
    return 0.0

def export_class_excel(class_id: int, date_str: Optional[str] = None, records: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    if not date_str:
        date_str = datetime.now().strftime("%Y-%m-%d")
        
    cls_list = get_classes()
    cls_info = next((c for c in cls_list if c["id"] == class_id), None)
    class_name = cls_info["class_name"] if cls_info else f"Class_{class_id}"
    
    attendance = records
    if not attendance:
        attendance = get_class_attendance_grades(class_id, date_str)
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Báo Cáo Lớp Học"
    
    ws.merge_cells("A1:I1")
    ws["A1"] = f"BÁO CÁO ĐIỂM DANH & ĐIỂM BÀI HỌC - {class_name.upper()} ({date_str})"
    ws["A1"].font = Font(size=14, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill(start_color="1E1B4B", end_color="1E1B4B", fill_type="solid")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36
    
    headers = ["STT", "Họ và Tên", "Điểm Danh", "Check 1", "Check 2", "BTVN", "BTVN - Check 2", "Check 2 - Check 1", "Cần Cố Gắng (Dưới TB)"]
    ws.cell(row=2, column=1, value="")
    ws.row_dimensions[3].height = 26
    for col_idx, h in enumerate(headers, 1):
        ws.cell(row=3, column=col_idx, value=h)
    
    header_fill = PatternFill(start_color="312E81", end_color="312E81", fill_type="solid")
    header_font = Font(name="Times New Roman", color="FFFFFF", bold=True, size=13)
    data_font = Font(name="Times New Roman", size=13)
    thin_border = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    for col_num in range(1, 10):
        cell = ws.cell(row=3, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border

    start_row = 4
    end_row = start_row + len(attendance) - 1 if len(attendance) > 0 else start_row
    avg_row_idx = end_row + 1

    for idx, r in enumerate(attendance, 1):
        curr_row = start_row + idx - 1
        st_name = str(r.get("student_name", ""))
        status_val = str(r.get("status", "Có mặt"))
        c1 = clean_num(r.get("check_1"))
        c2 = clean_num(r.get("check_2"))
        hw = clean_num(r.get("homework"))

        ws.cell(row=curr_row, column=1, value=f"=ROW()-3")
        ws.cell(row=curr_row, column=2, value=st_name)
        ws.cell(row=curr_row, column=3, value=status_val)
        
        c1_cell = ws.cell(row=curr_row, column=4, value=c1)
        c2_cell = ws.cell(row=curr_row, column=5, value=c2)
        hw_cell = ws.cell(row=curr_row, column=6, value=hw)
        
        c1_cell.number_format = '0.0'
        c2_cell.number_format = '0.0'
        hw_cell.number_format = '0.0'

        c7_cell = ws.cell(row=curr_row, column=7, value=f"=ROUNDUP(ABS(F{curr_row}-E{curr_row}), 1)")
        c8_cell = ws.cell(row=curr_row, column=8, value=f"=ROUNDUP(ABS(E{curr_row}-D{curr_row}), 1)")
        c7_cell.number_format = '0.0'
        c8_cell.number_format = '0.0'

        ws.cell(
            row=curr_row,
            column=9,
            value=f'=IF(C{curr_row}="Vắng mặt", "Vắng mặt", IF(_xlfn.TEXTJOIN(", ", TRUE, IF(AND(D{curr_row}>0, D{curr_row}<D${avg_row_idx}), "Check 1", ""), IF(AND(E{curr_row}>0, E{curr_row}<E${avg_row_idx}), "Check 2", ""), IF(AND(F{curr_row}>0, F{curr_row}<F${avg_row_idx}), "BTVN", ""))="", "Đạt yêu cầu", "⚠️ Cần cố gắng (" & _xlfn.TEXTJOIN(", ", TRUE, IF(AND(D{curr_row}>0, D{curr_row}<D${avg_row_idx}), "Check 1", ""), IF(AND(E{curr_row}>0, E{curr_row}<E${avg_row_idx}), "Check 2", ""), IF(AND(F{curr_row}>0, F{curr_row}<F${avg_row_idx}), "BTVN", "")) & ")"))'
        )

        for col_num in range(1, 10):
            c_cell = ws.cell(row=curr_row, column=col_num)
            c_cell.font = data_font
            c_cell.border = thin_border
            c_cell.alignment = Alignment(horizontal="center", vertical="center")

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if len(attendance) > 0:
        table_ref = f"A3:I{end_row}"
        tab = Table(displayName=f"ClassTable_{ts}", ref=table_ref)
        tab.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False
        )
        ws.add_table(tab)

        red_font = Font(color="B91C1C", bold=True)
        green_font = Font(color="15803D", bold=True)
        grey_font = Font(color="64748B", bold=True)

        rule_red = FormulaRule(
            formula=['NOT(ISERROR(SEARCH("Cần cố gắng", I4)))'],
            font=red_font
        )
        rule_green = FormulaRule(
            formula=['NOT(ISERROR(SEARCH("Đạt yêu cầu", I4)))'],
            font=green_font
        )
        rule_grey = FormulaRule(
            formula=['NOT(ISERROR(SEARCH("Vắng mặt", I4)))'],
            font=grey_font
        )
        ws.conditional_formatting.add(f"I4:I{end_row}", rule_red)
        ws.conditional_formatting.add(f"I4:I{end_row}", rule_green)
        ws.conditional_formatting.add(f"I4:I{end_row}", rule_grey)

    # Average row
    ws.cell(row=avg_row_idx, column=1, value="")
    ws.cell(row=avg_row_idx, column=2, value="Điểm trung bình (Average)")
    ws.cell(row=avg_row_idx, column=3, value="")

    if len(attendance) > 0:
        c1_vals = [clean_num(r.get("check_1")) for r in attendance if clean_num(r.get("check_1")) > 0 and str(r.get("status")) != "Vắng mặt"]
        c2_vals = [clean_num(r.get("check_2")) for r in attendance if clean_num(r.get("check_2")) > 0 and str(r.get("status")) != "Vắng mặt"]
        hw_vals = [clean_num(r.get("homework")) for r in attendance if clean_num(r.get("homework")) > 0 and str(r.get("status")) != "Vắng mặt"]
        
        avg_1 = trunc_1_dec(sum(c1_vals) / len(c1_vals)) if c1_vals else 0.0
        avg_2 = trunc_1_dec(sum(c2_vals) / len(c2_vals)) if c2_vals else 0.0
        avg_hw = trunc_1_dec(sum(hw_vals) / len(hw_vals)) if hw_vals else 0.0

        c1_avg_cell = ws.cell(row=avg_row_idx, column=4, value=avg_1)
        c2_avg_cell = ws.cell(row=avg_row_idx, column=5, value=avg_2)
        hw_avg_cell = ws.cell(row=avg_row_idx, column=6, value=avg_hw)

        c1_avg_cell.number_format = '0.0'
        c2_avg_cell.number_format = '0.0'
        hw_avg_cell.number_format = '0.0'

        diff_hw_c2 = trunc_1_dec(abs(avg_hw - avg_2))
        diff_c2_c1 = trunc_1_dec(abs(avg_2 - avg_1))

        c7_avg_cell = ws.cell(row=avg_row_idx, column=7, value=diff_hw_c2)
        c8_avg_cell = ws.cell(row=avg_row_idx, column=8, value=diff_c2_c1)

        c7_avg_cell.number_format = '0.0'
        c8_avg_cell.number_format = '0.0'

        ws.cell(row=avg_row_idx, column=9, value=f"=IF(D{avg_row_idx}>0, \"Đã tính TB lớp\", \"Chưa đủ điểm\")")

    avg_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    avg_font = Font(name="Times New Roman", bold=True, size=13, color="92400E")
    for col_num in range(1, 10):
        c_cell = ws.cell(row=avg_row_idx, column=col_num)
        c_cell.font = avg_font
        c_cell.fill = avg_fill
        c_cell.border = thin_border
        c_cell.alignment = Alignment(horizontal="center", vertical="center")

    summary_labels = [
        ("Check 1", 4, "D"),
        ("Check 2", 5, "E"),
        ("Bài tập về nhà", 6, "F")
    ]

    for idx, (m_label, col_num, col_let) in enumerate(summary_labels):
        r_idx = avg_row_idx + 2 + idx
        ws.cell(row=r_idx, column=1, value="")
        sum_title = ws.cell(
            row=r_idx,
            column=2,
            value=(
                f'="Check 1 dưới TB (< " & TEXT({col_let}{avg_row_idx}, "0.0") & ")"' if m_label == "Check 1" else
                f'="Check 2 dưới TB (< " & TEXT({col_let}{avg_row_idx}, "0.0") & ")"' if m_label == "Check 2" else
                f'="BTVN dưới TB (< " & TEXT({col_let}{avg_row_idx}, "0.0") & ")"'
            )
        )
        sum_title.font = Font(name="Times New Roman", bold=True, color="7F1D1D", size=13)
        sum_title.alignment = Alignment(horizontal="left", vertical="center")
        ws.merge_cells(f"C{r_idx}:I{r_idx}")
        val_cell = ws.cell(
            row=r_idx,
            column=3,
            value=f'=_xlfn.TEXTJOIN(", ", TRUE, _xlfn.FILTER(B{start_row}:B{end_row}, ({col_let}{start_row}:{col_let}{end_row}>0)*({col_let}{start_row}:{col_let}{end_row}<{col_let}{avg_row_idx})*(C{start_row}:C{end_row}<>"Vắng mặt"), "Không có (Tất cả đạt)"))' if len(attendance) > 0 else "Không có (Tất cả đạt)"
        )
        val_cell.font = Font(name="Times New Roman", bold=True, color="1E1E2F", size=13)
        val_cell.alignment = Alignment(horizontal="left", vertical="center")

    total_max_row = avg_row_idx + 5
    for col_idx in range(1, 10):
        col_let = openpyxl.utils.get_column_letter(col_idx)
        max_len = 0
        for r_idx in range(3, total_max_row + 1):
            if r_idx > end_row and col_idx == 3:
                continue
            cell_val = ws.cell(row=r_idx, column=col_idx).value
            val_str = str(cell_val) if cell_val is not None else ""
            if val_str.startswith("="):
                if col_idx == 9:
                    val_str = "⚠️ Cần cố gắng (Check 1, Check 2, BTVN)"
                elif col_idx == 2:
                    val_str = "Check 1 dưới TB (< 10.0)"
                elif col_idx == 1:
                    val_str = "999"
                elif col_idx in (7, 8):
                    val_str = "10.0"
                else:
                    val_str = ""
            if len(val_str) > max_len:
                max_len = len(val_str)

        extra_padding = 12 if col_idx == 9 else (8 if col_idx == 2 else 5)
        min_w = 54 if col_idx == 9 else (36 if col_idx == 2 else (16 if col_idx == 3 else 14))
        ws.column_dimensions[col_let].width = max(max_len + extra_padding, min_w)

    files_dir = get_setting("files_dir")
    os.makedirs(files_dir, exist_ok=True)
    filename = f"ClassReport_{class_name}_{date_str}_{ts}.xlsx"
    filepath = os.path.join(files_dir, filename)
    wb.save(filepath)
    return {"status": "success", "filename": filename, "filepath": filepath}

def export_class_docx(class_id: int, date_str: Optional[str] = None, records: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    if not date_str:
        date_str = datetime.now().strftime("%Y-%m-%d")
        
    cls_list = get_classes()
    cls_info = next((c for c in cls_list if c["id"] == class_id), None)
    class_name = cls_info["class_name"] if cls_info else f"Class_{class_id}"
    
    attendance = records
    if not attendance:
        attendance = get_class_attendance_grades(class_id, date_str)
    
    doc = docx.Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"BÁO CÁO NGHỈ HỌC & ĐIỂM BÀI HỌC\nLỚP: {class_name.upper()} - NGÀY: {date_str}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 27, 75)
    
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["STT", "Họ và Tên", "Điểm Danh", "Check 1", "Check 2", "BTVN"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    for idx, r in enumerate(attendance, 1):
        c1 = r.get("check_1", 0)
        c2 = r.get("check_2", 0)
        hw = r.get("homework", 0)
        hw_str = "Không BTVN" if hw == 0 else str(hw)
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = r.get("student_name", "")
        row_cells[2].text = r.get("status", "")
        row_cells[3].text = str(c1)
        row_cells[4].text = str(c2)
        row_cells[5].text = hw_str
        
    files_dir = get_setting("files_dir")
    os.makedirs(files_dir, exist_ok=True)
    ts = datetime.now().strftime("%H%M%S")
    filename = f"ClassReport_{class_name}_{date_str}_{ts}.docx"
    filepath = os.path.join(files_dir, filename)
    doc.save(filepath)
    return {"status": "success", "filename": filename, "filepath": filepath}
