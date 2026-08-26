import os
import re
import json
from typing import List, Dict, Any
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

from services.compiler import get_instruction_map

OPTION_PREFIX_RE = re.compile(r'^\s*(?:\*\*)?\s*([A-E])\s*(?:\.(?:\*\*)?|(?:\*\*)?\.|(?:\*\*)?\))\s*(.*)', re.DOTALL)
QUESTION_RE = re.compile(r'^\s*(?:\*\*)?(?:Question|Câu|Q|Sentence)\s*(\d+)\s*[:.\-]?\s*(?:\*\*)?\s*[:.\-]?\s*(.*)', re.IGNORECASE | re.DOTALL)

def normalize_q_key(q_val):
    try:
        return int(q_val)
    except (ValueError, TypeError):
        match = re.search(r'\d+', str(q_val))
        if match:
            return int(match.group(0))
    return str(q_val)

def is_run_yellow_or_highlighted(run) -> bool:
    """Checks if a docx Run has yellow highlight, general highlight, yellow shading, or colored emphasis."""
    try:
        if run.font.highlight_color is not None and run.font.highlight_color != WD_COLOR_INDEX.AUTO:
            return True
    except Exception:
        pass

    try:
        r_elem = run._r
        rPr = r_elem.rPr
        if rPr is not None:
            highlights = rPr.xpath('./w:highlight')
            if highlights:
                val = (highlights[0].get(qn('w:val')) or highlights[0].get('val') or '').lower()
                if val and val != 'none':
                    return True

            shadings = rPr.xpath('./w:shd')
            if shadings:
                fill = (shadings[0].get(qn('w:fill')) or shadings[0].get('fill') or '').upper()
                if fill and fill not in ('AUTO', 'NONE', 'FFFFFF', '000000', 'CLEAR'):
                    return True

            colors = rPr.xpath('./w:color')
            if colors:
                val = (colors[0].get(qn('w:val')) or colors[0].get('val') or '').upper()
                if val in ('FFFF00', 'FFD700', 'FFCC00', 'FFE600', 'E6B800', 'FF0000', '00B050', '385723', 'ED7D31'):
                    return True
    except Exception:
        pass

    try:
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            if rgb == RGBColor(255, 0, 0) or (rgb[0] > 180 and rgb[1] > 180 and rgb[2] < 120):
                return True
    except Exception:
        pass

    return False

def split_paragraph_runs_by_tab(paragraph):
    options_runs = []
    current_runs = []
    
    for run in paragraph.runs:
        is_highlighted = is_run_yellow_or_highlighted(run)
        if '\t' in run.text:
            parts = run.text.split('\t')
            if parts[0]:
                current_runs.append({
                    'text': parts[0],
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'highlight': run.font.highlight_color,
                    'color_rgb': run.font.color.rgb if run.font.color else None,
                    'is_yellow': is_highlighted
                })
            options_runs.append(current_runs)
            
            for part in parts[1:-1]:
                options_runs.append([{
                    'text': part,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'highlight': run.font.highlight_color,
                    'color_rgb': run.font.color.rgb if run.font.color else None,
                    'is_yellow': is_highlighted
                }] if part else [])
                
            current_runs = []
            if parts[-1]:
                current_runs.append({
                    'text': parts[-1],
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'highlight': run.font.highlight_color,
                    'color_rgb': run.font.color.rgb if run.font.color else None,
                    'is_yellow': is_highlighted
                })
        else:
            current_runs.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'highlight': run.font.highlight_color,
                'color_rgb': run.font.color.rgb if run.font.color else None,
                'is_yellow': is_highlighted
            })
            
    if current_runs:
        options_runs.append(current_runs)
        
    # If not split by tabs, check if multiple options exist in a single line e.g. A. ... B. ... C. ... D. ...
    if len(options_runs) == 1 and options_runs[0]:
        combined_text = "".join(r['text'] for r in options_runs[0])
        matches = list(re.finditer(r'(?:\s{2,}|\b|^)(?:[A-E]\.|\([A-E]\)|[A-E]\))\s+', combined_text))
        if len(matches) > 1:
            split_opts = []
            flat_runs = options_runs[0]
            for idx, m in enumerate(matches):
                start_char = m.start()
                end_char = matches[idx + 1].start() if idx + 1 < len(matches) else len(combined_text)
                
                curr_char = 0
                opt_chunk = []
                for r in flat_runs:
                    r_len = len(r['text'])
                    r_start = curr_char
                    r_end = curr_char + r_len
                    if r_end > start_char and r_start < end_char:
                        overlap_start = max(0, start_char - r_start)
                        overlap_end = min(r_len, end_char - r_start)
                        opt_chunk.append({
                            'text': r['text'][overlap_start:overlap_end],
                            'bold': r['bold'],
                            'italic': r['italic'],
                            'underline': r['underline'],
                            'highlight': r['highlight'],
                            'color_rgb': r['color_rgb'],
                            'is_yellow': r['is_yellow']
                        })
                    curr_char = r_end
                if opt_chunk:
                    split_opts.append(opt_chunk)
            if split_opts:
                return split_opts

    return options_runs

def run_dicts_to_markdown(run_dicts):
    md_parts = []
    i = 0
    while i < len(run_dicts):
        run = run_dicts[i]
        text = run['text']
        if not text:
            i += 1
            continue
        
        if run['underline'] and (i + 1 < len(run_dicts)):
            next_run = run_dicts[i + 1]
            if next_run['bold'] and next_run['text'].strip() in ['(A)', '(B)', '(C)', '(D)', '(E)']:
                let = next_run['text'].strip().replace('(', '').replace(')', '')
                clean_text = text.strip()
                leading = text[:len(text)-len(text.lstrip())]
                trailing = text[len(text.rstrip()):]
                md_parts.append(f"{leading}[{clean_text}]({let}){trailing}")
                i += 2
                continue
                
        bold = run['bold']
        italic = run['italic']
        underline = run['underline']
        
        clean_text = text.strip()
        if not clean_text:
            md_parts.append(text)
            i += 1
            continue
            
        leading = text[:len(text)-len(text.lstrip())]
        trailing = text[len(text.rstrip()):]
        
        formatted = clean_text
        if underline:
            formatted = f"[{formatted}]"
        if bold:
            formatted = f"**{formatted}**"
        if italic:
            formatted = f"*{formatted}*"
            
        md_parts.append(f"{leading}{formatted}{trailing}")
        i += 1
    return "".join(md_parts)

def parse_runs_to_markdown(paragraph) -> str:
    run_dicts = []
    for run in paragraph.runs:
        run_dicts.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'highlight': run.font.highlight_color,
            'color_rgb': run.font.color.rgb if run.font.color else None,
            'is_yellow': is_run_yellow_or_highlighted(run)
        })
    return run_dicts_to_markdown(run_dicts)

def parse_answer_key_table(doc) -> Dict[int, str]:
    """Comprehensively parses answer keys from tables and bottom answer key sections in DOCX."""
    answers: Dict[int, str] = {}
    
    # 1. Check all tables in the document
    for table in doc.tables:
        rows = table.rows
        num_rows = len(rows)
        if num_rows == 0:
            continue
            
        # Format A: Check individual cells with '1.A', '1. A', '1 - A', '1A', 'Câu 1: A'
        for row in rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                matches = re.findall(r'(?:(?:Câu|Q|Question)\s*)?(\d+)\s*[:.\-]?\s*([A-E])\b', text, re.IGNORECASE)
                for q_str, a_str in matches:
                    answers[int(q_str)] = a_str.upper()
                    
        # Format B: Matrix table with Q numbers in row i and answers in row i+1
        for r_idx in range(num_rows - 1):
            row_q = [c.text.strip() for c in rows[r_idx].cells]
            row_a = [c.text.strip() for c in rows[r_idx + 1].cells]
            if len(row_q) == len(row_a) and len(row_q) >= 2:
                for q_cell, a_cell in zip(row_q, row_a):
                    q_m = re.match(r'^(?:(?:Câu|Q)\s*)?(\d+)$', q_cell, re.IGNORECASE)
                    a_m = re.match(r'^([A-E])$', a_cell, re.IGNORECASE)
                    if q_m and a_m:
                        answers[int(q_m.group(1))] = a_m.group(1).upper()

    # 2. Check paragraphs for dedicated Answer Key sections (e.g. at the bottom of the document)
    is_in_answer_section = False
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not p_text:
            continue
            
        header_match = re.search(r'(?:BẢNG\s+)?ĐÁP\s*ÁN|ANSWER\s*KEY|HƯỚNG\s*DẪN\s*CHẤM|KEY\s*ĐÁP\s*ÁN', p_text, re.IGNORECASE)
        if header_match:
            is_in_answer_section = True
            
        if is_in_answer_section:
            matches = re.findall(r'(?:(?:Câu|Q|Question)\s*)?(\d+)\s*[:.\-]?\s*([A-E])\b', p_text, re.IGNORECASE)
            for q_str, a_str in matches:
                answers[int(q_str)] = a_str.upper()
                
    return answers

def match_instruction(text: str) -> str:
    text_lower = text.lower()
    inst_map = get_instruction_map()
    
    for k, v in inst_map.items():
        if text.strip() == v.strip():
            return k
            
    if "pronunciation" in text_lower and "differ" in text_lower:
        return "pr"
    if "stress" in text_lower and "differ" in text_lower:
        return "st"
    if "closest in meaning" in text_lower:
        return "sy"
    if "opposite in meaning" in text_lower:
        return "an"
    if "meaning of the sign" in text_lower:
        return "sg"
    if "meaning of the notice" in text_lower:
        return "nt"
    if "fits each of the numbered blanks" in text_lower:
        return "cz"
    if "arrangement of the sentences" in text_lower:
        return "ro"
    if "read the following passage" in text_lower:
        return "rd"
    if "needs correction" in text_lower or "needs-correction" in text_lower or "underlined part that needs correction" in text_lower:
        return "er"
    if "correct answer" in text_lower:
        return "mq"
        
    return None

def iter_blocks(doc):
    for child in doc.element.body:
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)

def convert_docx_to_json(filepath: str) -> List[Dict[str, Any]]:
    doc = Document(filepath)
    answers_map = parse_answer_key_table(doc)
    
    exercises = []
    active_type = None
    
    passage_paragraphs = []
    reorder_sentences = []
    notice_body = ""
    current_question = None
    current_sub_question = None
    last_block_was_question_or_option = False
    
    parent_exercise = None
    is_first_table = True
    
    for block in iter_blocks(doc):
        if isinstance(block, Table):
            if is_first_table:
                is_first_table = False
                continue
            continue
            
        p_text = block.text.strip()
        if not p_text:
            continue
            
        inst_type = match_instruction(p_text)
        if inst_type:
            active_type = inst_type
            current_instruction = p_text
            passage_paragraphs = []
            reorder_sentences = []
            notice_body = ""
            current_question = None
            current_sub_question = None
            parent_exercise = None
            last_block_was_question_or_option = False
            continue
            
        if p_text.upper().startswith("UNIT") and active_type is None:
            continue
            
        p_markdown = parse_runs_to_markdown(block)
        
        q_match = QUESTION_RE.match(p_markdown)
        if q_match:
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()
            
            last_block_was_question_or_option = True
            
            if active_type in ["cz", "rd"]:
                if parent_exercise is None or passage_paragraphs:
                    parent_exercise = {
                        "t": active_type,
                        "instruction": current_instruction if 'current_instruction' in locals() else "",
                        "b": passage_paragraphs.copy(),
                        "k": []
                    }
                    exercises.append(parent_exercise)
                    passage_paragraphs.clear()
                    
                current_sub_question = {
                    "q": q_num,
                    "x": q_text,
                    "instruction": current_instruction if 'current_instruction' in locals() else "",
                    "o": [],
                    "a": ""
                }
                parent_exercise["k"].append(current_sub_question)
                current_question = None
                
            elif active_type == "ro":
                current_question = {
                    "t": "ro",
                    "q": q_num,
                    "x": q_text,
                    "instruction": current_instruction if 'current_instruction' in locals() else "",
                    "i": reorder_sentences.copy(),
                    "o": [],
                    "a": ""
                }
                exercises.append(current_question)
                reorder_sentences.clear()
                current_sub_question = None
                
            elif active_type == "nt":
                current_question = {
                    "t": "nt",
                    "q": q_num,
                    "x": q_text,
                    "instruction": current_instruction if 'current_instruction' in locals() else "",
                    "b": "",
                    "o": [],
                    "a": ""
                }
                exercises.append(current_question)
                current_sub_question = None
                
            else:
                current_question = {
                    "t": active_type or "mq",
                    "q": q_num,
                    "x": q_text,
                    "instruction": current_instruction if 'current_instruction' in locals() else "",
                    "o": [],
                    "a": ""
                }
                exercises.append(current_question)
                current_sub_question = None
                
            continue
            
        options_split_runs = split_paragraph_runs_by_tab(block)
        is_option_paragraph = False
        parsed_options = []
        
        for option_runs in options_split_runs:
            opt_md = run_dicts_to_markdown(option_runs)
            opt_match = OPTION_PREFIX_RE.match(opt_md)
            if opt_match:
                is_option_paragraph = True
                opt_letter = opt_match.group(1).upper()
                opt_val = opt_match.group(2).strip()
                
                is_correct = any(r.get('is_yellow') or r.get('highlight') is not None or r.get('color_rgb') == RGBColor(255, 0, 0) for r in option_runs)
                parsed_options.append((opt_letter, opt_val, is_correct))
                
        if is_option_paragraph:
            last_block_was_question_or_option = True
            target = current_sub_question if active_type in ["cz", "rd"] else current_question
            
            if target:
                for letter, val, is_correct in parsed_options:
                    target["o"].append(val)
                    if is_correct:
                        target["a"] = letter
            continue
            
        if active_type == "nt" and current_question and not current_question["b"]:
            notice_txt = p_markdown.strip()
            if notice_txt.startswith("*") and notice_txt.endswith("*"):
                notice_txt = notice_txt[1:-1].strip()
            current_question["b"] = notice_txt
            last_block_was_question_or_option = False
            
        elif active_type == "ro":
            reorder_sentences.append(p_markdown)
            last_block_was_question_or_option = False
            
        elif active_type in ["cz", "rd"]:
            if last_block_was_question_or_option:
                passage_paragraphs.clear()
                parent_exercise = None
                
            passage_paragraphs.append(p_markdown)
            last_block_was_question_or_option = False
            
    for ex in exercises:
        ex_type = ex.get("t")
        if ex_type in ["cz", "rd"]:
            for sub in ex.get("k", []):
                q_key = normalize_q_key(sub.get("q"))
                if q_key in answers_map and not sub.get("a"):
                    sub["a"] = answers_map[q_key]
        else:
            q_key = normalize_q_key(ex.get("q"))
            if q_key in answers_map and not ex.get("a"):
                ex["a"] = answers_map[q_key]
                
    return exercises
